"""Insert the Hepatic Steatosis (S0–S3) dataset content into the existing
SmartLiva progress report.

The source .docx is used verbatim as the template: nothing that is already in
the document is renamed, restyled or renumbered.  Every new paragraph is a deep
copy of a paragraph that already exists in the file (body / bullet / sub-heading),
so the inserted block inherits the original TH SarabunPSK formatting exactly.

Figures are the ones counted directly from the Zenodo .mat file
(DOI 10.5281/zenodo.1009146, MD5 c87da28a498eae0f0874408c7ac92524).
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

SOURCE = Path("/Users/king_phuripol/Downloads/SmartLiva_Ultrasound_AI_Training_Report(1).docx")
OUTPUT = Path(
    "/Users/king_phuripol/AI-Engineer/01_Projects/Aong-Task/SmartMed/"
    "New-SmartLiva/reports/SmartLiva_Ultrasound_AI_Training_Report_Updated.docx"
)

# Anchor texts (matched exactly against the untouched source document).
ANCHOR_SUBHEAD = "5.5 การประเมินภาวะไขมันพอกตับและความเสี่ยงพยาธิใบไม้ตับ (Steatosis & CCA Risk Specialists)"
ANCHOR_BULLET = "Liver Fluke & CCA Risk Specialist:"
ANCHOR_BODY = "การประเมินประสิทธิภาพของระบบ SmartLiva"

GRADE_ROWS = (
    ("ระดับความรุนแรง (เกณฑ์ Biopsy Fat%)", "จำนวนผู้ป่วย (ราย)", "ภาพทั้งหมด (เฟรม)", "ภาพไม่ซ้ำ (Unique)"),
    ("S0 – ไม่มีไขมันพอกตับ (น้อยกว่า 5%)", "17", "170", "168"),
    ("S1 – ระดับเริ่มต้น (5–33%)", "20", "200", "200"),
    ("S2 – ระดับปานกลาง (มากกว่า 33–66%)", "8", "80", "76"),
    ("S3 – ระดับรุนแรง (มากกว่า 66%)", "10", "100", "100"),
    ("รวมทั้งหมด", "55", "550", "544"),
)
COL_WIDTHS = (3600, 1750, 1950, 2050)  # dxa, sums to the 9350 used by the report body


def find_paragraph(document, needle, *, exact=True):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if (text == needle) if exact else text.startswith(needle):
            return paragraph
    raise LookupError(needle)


def clone(template_paragraph, text):
    """Copy a paragraph verbatim, keep a single run and swap in new text."""
    element = deepcopy(template_paragraph._p)
    runs = element.findall(qn("w:r"))
    for extra in runs[1:]:
        element.remove(extra)
    for t in runs[0].findall(qn("w:t")):
        runs[0].remove(t)
    t = runs[0].makeelement(qn("w:t"), {})
    t.set(qn("xml:space"), "preserve")
    t.text = text
    runs[0].append(t)
    return element


def build_table(document, template_table, template_run_rpr):
    """A table that matches the report's existing TableGrid tables."""
    table = document.add_table(rows=len(GRADE_ROWS), cols=4)
    table._tbl.remove(table._tbl.tblPr)
    table._tbl.insert(0, deepcopy(template_table._tbl.tblPr))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in COL_WIDTHS:
        col = grid.makeelement(qn("w:gridCol"), {})
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row, values in zip(table.rows, GRADE_ROWS):
        for cell, width, value in zip(row.cells, COL_WIDTHS, values):
            tc_pr = cell._tc.get_or_add_tcPr()
            for existing in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(existing)
            tc_w = tc_pr.makeelement(qn("w:tcW"), {})
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            run._r.insert(0, deepcopy(template_run_rpr))

    # Keep the whole table on one page and repeat the header if it ever wraps.
    for index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(tr_pr.makeelement(qn("w:cantSplit"), {}))
        if index == 0:
            header = tr_pr.makeelement(qn("w:tblHeader"), {})
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        if index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    return table


def main():
    doc = Document(SOURCE)

    tpl_subhead = find_paragraph(doc, ANCHOR_SUBHEAD)
    tpl_bullet = find_paragraph(doc, ANCHOR_BULLET, exact=False)
    tpl_body = find_paragraph(doc, ANCHOR_BODY, exact=False)
    run_rpr = tpl_body._p.findall(qn("w:r"))[0].find(qn("w:rPr"))

    blocks = []
    blocks.append(clone(tpl_subhead, "ชุดข้อมูลอ้างอิงสำหรับพัฒนาโมเดลไขมันพอกตับ (Hepatic Steatosis Dataset: S0–S3)"))
    blocks.append(clone(
        tpl_body,
        "เพื่อยกระดับโมดูลไขมันพอกตับจากการวิเคราะห์เชิงกฎ (Rule-based Acoustic Analysis) ไปสู่โมเดลที่เรียนรู้จากผลชิ้นเนื้อจริง "
        "ทีมงานได้จัดหาและตรวจสอบชุดข้อมูลอ้างอิงระดับสากลที่มี Ground Truth เป็นผลตรวจ Liver Biopsy รายผู้ป่วย ดังรายละเอียดต่อไปนี้:",
    ))
    blocks.append(clone(
        tpl_bullet,
        "ชื่อชุดข้อมูลและแหล่งเผยแพร่: Dataset of B-mode fatty liver ultrasound images (Byra M. et al., IJCARS 2018) "
        "เผยแพร่บน Zenodo DOI 10.5281/zenodo.1009146 ไฟล์ dataset_liver_bmodes_steatosis_assessment_IJCARS.mat "
        "ตรวจสอบค่า MD5 ได้ c87da28a498eae0f0874408c7ac92524 ตรงกับระเบียนต้นทาง",
    ))
    blocks.append(clone(
        tpl_bullet,
        "Ground Truth และประชากร: ใช้ค่าเปอร์เซ็นต์เซลล์ตับที่มีไขมันสะสม (% Hepatocytes with Steatosis) จาก Wedge Liver Biopsy "
        "รายผู้ป่วย ช่วง 0–85% เก็บภาพจากผู้ป่วยโรคอ้วนรุนแรง 55 รายในประเทศโปแลนด์ ด้วยเครื่อง GE Vivid E9 หัวตรวจ Sector 2.5 MHz "
        "จึงไม่ใช่ประชากรไทยและต้องมี External Validation กับผู้ป่วยไทยก่อนใช้งานทางคลินิก",
    ))
    blocks.append(clone(
        tpl_bullet,
        "โครงสร้างข้อมูล: ผู้ป่วยแต่ละรายมี 10 เฟรมต่อเนื่อง ขนาดภาพ 434 x 636 พิกเซล แบบ Grayscale 8-bit เมื่อตรวจสอบด้วยค่า Hash "
        "พบภาพซ้ำตรงกัน 6 ภาพ ซึ่งทั้งหมดเป็นการซ้ำภายในผู้ป่วยคนเดียวกัน ไม่ข้ามผู้ป่วย จึงไม่ขัดกับกฎ Zero Cross-Patient Leakage ในข้อ 3",
    ))
    intro = clone(tpl_body, "จำนวนผู้ป่วยและภาพจำแนกตามระดับความรุนแรง S0–S3 (ตรวจนับจากไฟล์ต้นฉบับโดยตรง):")
    p_pr = intro.find(qn("w:pPr"))
    p_pr.insert(0, p_pr.makeelement(qn("w:keepNext"), {}))
    blocks.append(intro)

    table = build_table(doc, doc.tables[3], run_rpr)
    blocks.append(table._tbl)

    blocks.append(clone(
        tpl_bullet,
        "เกณฑ์การแปลงระดับความรุนแรง: S0 น้อยกว่า 5%, S1 ระหว่าง 5–33%, S2 มากกว่า 33% ถึง 66% และ S3 มากกว่า 66% "
        "ทั้งนี้ไฟล์ต้นฉบับกำหนดสถานะ Positive เมื่อค่าไขมันตั้งแต่ 5% ขึ้นไป ผู้ป่วยที่มีค่าไขมัน 5% พอดีจำนวน 1 รายจึงถูกจัดอยู่ในกลุ่ม S1",
    ))
    blocks.append(clone(
        tpl_bullet,
        "การแบ่งข้อมูลสำหรับฝึกสอน: 10 เฟรมจากผู้ป่วยคนเดียวกันไม่ถือเป็นตัวอย่างอิสระทางสถิติ การแบ่ง Train / Validation / Test "
        "จึงต้องกระทำที่ระดับผู้ป่วย (Patient-level Split) เท่านั้น สอดคล้องกับกฎเหล็ก Zero Cross-Patient Leakage ของโครงการ",
    ))
    blocks.append(clone(
        tpl_bullet,
        "สถานะการใช้งาน: ตัวเลขข้างต้นหมายถึงชุดข้อมูลที่จัดเตรียมพร้อมสำหรับการฝึกสอน (Training-ready Dataset) "
        "โดยอยู่ระหว่างการพัฒนา Training Pipeline เพื่อนำมาใช้กับโมดูล src/workflow/specialists/fatty_liver.py ในระยะถัดไป",
    ))
    blocks.append(clone(
        tpl_body,
        "ช่องว่างด้านข้อมูลและเป้าหมายการจัดหาเพิ่มเติม (Data Gap Analysis): ชุดข้อมูลนี้เหมาะสำหรับเริ่มพัฒนาโมเดลด้วยเทคนิค "
        "Transfer Learning หรือการทำ Regression ค่าไขมันโดยตรง แต่จำนวนผู้ป่วยยังน้อยเกินไปสำหรับโมเดล S0–S3 ที่พร้อมใช้งานทางคลินิก "
        "โดยเฉพาะกลุ่ม S2 ที่มีเพียง 8 ราย และกลุ่ม S3 ที่มีเพียง 10 ราย ทีมงานจึงกำหนดเป้าหมายการจัดหาข้อมูลเพิ่มเติมดังนี้:",
    ))
    blocks.append(clone(
        tpl_bullet,
        "กรณีเป้าหมาย 1,000 ผู้ป่วยครอบคลุมทุกระดับ S0–S3: ต้องจัดหาเพิ่มอีก 945 ผู้ป่วย ซึ่งเทียบเท่าประมาณ 9,450 เฟรม "
        "หากเก็บ 10 เฟรมต่อผู้ป่วยตาม Protocol เดิม",
    ))
    blocks.append(clone(
        tpl_bullet,
        "กรณีเป้าหมาย 1,000 ผู้ป่วยเฉพาะกลุ่มที่มีภาวะไขมันพอกตับ (S1–S3): ต้องจัดหาเพิ่มอีก 962 ผู้ป่วย เนื่องจากปัจจุบันมีผู้ป่วยกลุ่ม "
        "S1–S3 เพียง 38 ราย (380 เฟรม) และกลุ่มควบคุม S0 อีก 17 ราย (170 เฟรม)",
    ))
    blocks.append(clone(
        tpl_bullet,
        "ข้อควรระวังในการนับ: หน่วยนับที่ใช้กำหนดเป้าหมายต้องเป็นจำนวนผู้ป่วย ไม่ใช่จำนวนภาพ เพราะหลายเฟรมจากผู้ป่วยรายเดียวกัน "
        "ไม่ใช่ตัวอย่างอิสระ การรายงานเป็นจำนวนภาพจะทำให้ประเมินขนาดตัวอย่างสูงเกินจริงถึง 10 เท่า",
    ))
    blocks.append(clone(
        tpl_bullet,
        "ความจำเป็นด้าน External Validation: ต้องจัดหาชุดข้อมูลผู้ป่วยไทยจากโรงพยาบาลพันธมิตร (รพ.อุดรธานี) ที่มีผลยืนยันระดับ "
        "ไขมันพอกตับ เพื่อรองรับความแตกต่างด้านประชากร เครื่องตรวจ และ Protocol การสแกน ก่อนนำโมเดลไปใช้งานทางคลินิกจริง",
    ))

    # Append at the end of section 5.5, right before the section 6 heading.
    cursor = find_paragraph(doc, "6. เกณฑ์การวัดผลและผลการทดสอบความแม่นยำทางคลินิก (Evaluation & Benchmarks)")._p
    for element in blocks:
        cursor.addprevious(element)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
