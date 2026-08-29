"""Generate the Thai IP technical-disclosure document for SmartLiva in Word (.docx) format.

Answers the three questions raised by the intellectual-property reviewer:
  1. Structure, operation and technical characteristics of the newly designed learning
     models for lesion / fibrosis / cancer / liver-fluke detection.
  2. How the system ingests data for processing, risk assessment and lesion localisation,
     specifically where it departs from the prior work (SmartLiva-LiverUS-SDK v1.1.0).
  3. What the models have in common.

House style is copied verbatim from reports/SmartLiva_Dataset_Training_Report_TH_Updated.docx:
pure black-and-white, TH SarabunPSK only, US Letter, 1-inch margins, "หน้า N" footer.

Every figure quoted here was read out of the source tree or the training logs. Claims that
appear in README.md but are not implemented in code are deliberately excluded and are listed
instead in Appendix ค.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

THAI_FONT = "TH SarabunPSK"
# TH SarabunPSK has no arrow glyphs; connectors in the diagrams use a font that does.
ARROW_FONT = "Arial"
BLACK = RGBColor(0, 0, 0)

# Type scale, in points. Mirrors the half-point sizes found in the reference document.
SZ_TITLE = 26
SZ_SUBTITLE = 15
SZ_H1 = 18
SZ_H2 = 16
SZ_BODY = 14
SZ_SMALL = 13
SZ_TABLE = 12.5

CONTENT_WIDTH = 6.5  # inches: US Letter minus two 1-inch margins


# ---------------------------------------------------------------------------
# Run-level helpers
# ---------------------------------------------------------------------------

def _run(p, text, size=SZ_BODY, bold=False, italic=False, font=THAI_FONT):
    """Add a run with the complex-script attributes Word needs for Thai.

    python-docx only writes w:sz / w:rFonts@ascii. Thai is a complex script, so Word reads
    w:szCs, w:bCs and w:rFonts@cs instead -- without them Thai renders at the default 10pt
    no matter what size was requested. Elements are inserted in schema order.
    """
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = BLACK

    rPr = r._element.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font)

    b = rPr.find(qn("w:b"))
    if b is not None:
        bCs = OxmlElement("w:bCs")
        bCs.set(qn("w:val"), "1" if bold else "0")
        b.addnext(bCs)

    i = rPr.find(qn("w:i"))
    if i is not None:
        iCs = OxmlElement("w:iCs")
        iCs.set(qn("w:val"), "1" if italic else "0")
        i.addnext(iCs)

    sz = rPr.find(qn("w:sz"))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(round(size * 2))))
    if sz is not None:
        sz.addnext(szCs)
    else:
        rPr.append(szCs)
    return r


def _para(doc, space_before=0, space_after=6, line_spacing=1.15, align=None, indent=0.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if indent:
        pf.left_indent = Inches(indent)
    return p


# ---------------------------------------------------------------------------
# Block-level helpers
# ---------------------------------------------------------------------------

def h1(doc, text):
    p = _para(doc, space_before=16, space_after=8)
    p.paragraph_format.keep_with_next = True
    _run(p, text, size=SZ_H1, bold=True)
    return p


def h2(doc, text):
    p = _para(doc, space_before=12, space_after=5)
    p.paragraph_format.keep_with_next = True
    _run(p, text, size=SZ_H2, bold=True)
    return p


def body(doc, text, indent=0.0, space_after=6):
    p = _para(doc, space_after=space_after, indent=indent)
    _run(p, text, size=SZ_BODY)
    return p


def lead(doc, label, text, indent=0.0):
    """A paragraph whose opening phrase is bold."""
    p = _para(doc, space_after=6, indent=indent)
    _run(p, label, size=SZ_BODY, bold=True)
    _run(p, text, size=SZ_BODY)
    return p


def bullet(doc, text, level=0, label=""):
    p = _para(doc, space_after=4, indent=0.28 + 0.26 * level)
    pf = p.paragraph_format
    pf.first_line_indent = Inches(-0.20)
    marker = "•  " if level == 0 else "–  "
    _run(p, marker, size=SZ_BODY)
    if label:
        _run(p, label, size=SZ_BODY, bold=True)
    _run(p, text, size=SZ_BODY)
    return p


def numbered(doc, index, text, label="", indent=0.28):
    p = _para(doc, space_after=4, indent=indent)
    p.paragraph_format.first_line_indent = Inches(-0.28)
    _run(p, f"{index}. ", size=SZ_BODY, bold=True)
    if label:
        _run(p, label, size=SZ_BODY, bold=True)
    _run(p, text, size=SZ_BODY)
    return p


def caption(doc, text):
    p = _para(doc, space_before=4, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, text, size=SZ_SMALL, italic=True)
    return p


def page_break(doc):
    doc.add_page_break()


def rule(doc, space_before=6, space_after=10):
    """A horizontal line drawn as a paragraph bottom border."""
    p = _para(doc, space_before=space_before, space_after=space_after)
    pPr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    pPr.append(borders)
    return p


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _set_cell_borders(cell, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def _set_cell_margins(cell, top=90, bottom=90, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _repeat_as_header(row):
    """Mark a row so Word repeats it at the top of every page the table spans."""
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _no_split(row):
    """Keep a row's cells on one page instead of breaking mid-row."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _cell_text(cell, text, size=SZ_TABLE, bold=False, align=None, first=True):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    _run(p, text, size=size, bold=bold)
    return p


def table(doc, headers, rows, widths, size=SZ_TABLE, align_cols=None, space_after=10):
    """Black-ruled table. `widths` are inches and must sum to CONTENT_WIDTH."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    align_cols = align_cols or [None] * len(headers)

    _repeat_as_header(t.rows[0])
    _no_split(t.rows[0])
    for col, (label, width) in enumerate(zip(headers, widths)):
        cell = t.cell(0, col)
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell, sz="8")
        _set_cell_margins(cell)
        _cell_text(cell, label, size=size, bold=True, align=align_cols[col])

    for r, record in enumerate(rows, start=1):
        _no_split(t.rows[r])
        for col, value in enumerate(record):
            cell = t.cell(r, col)
            cell.width = Inches(widths[col])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            lines = str(value).split("\n")
            for k, line in enumerate(lines):
                _cell_text(cell, line, size=size, align=align_cols[col], first=(k == 0))

    tail = _para(doc, space_after=space_after)
    _run(tail, "", size=6)
    return t


# ---------------------------------------------------------------------------
# Diagrams -- drawn with Word tables so they stay vector, monochrome and editable
# ---------------------------------------------------------------------------

def dbox(doc, lines, width=4.6, size=SZ_TABLE, heavy=False):
    """One bordered box in a flow diagram. First line is bold."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _no_split(t.rows[0])
    cell = t.cell(0, 0)
    cell.width = Inches(width)
    _set_cell_borders(cell, sz="24" if heavy else "8")
    _set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    for k, line in enumerate(lines):
        _cell_text(cell, line, size=size, bold=(k == 0),
                   align=WD_ALIGN_PARAGRAPH.CENTER, first=(k == 0))
    return t


def drow(doc, columns, widths, size=SZ_TABLE):
    """A row of bordered boxes side by side. Each column is a list of lines."""
    t = doc.add_table(rows=1, cols=len(columns))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _no_split(t.rows[0])
    for col, (lines, width) in enumerate(zip(columns, widths)):
        cell = t.cell(0, col)
        cell.width = Inches(width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_borders(cell)
        _set_cell_margins(cell, top=80, bottom=80, left=110, right=110)
        for k, line in enumerate(lines):
            _cell_text(cell, line, size=size, bold=(k == 0),
                       align=WD_ALIGN_PARAGRAPH.CENTER, first=(k == 0))
    return t


def dgap(doc, height=4):
    p = _para(doc, space_before=0, space_after=0)
    _run(p, "", size=height)
    return p


def darrow(doc, label="", glyph="↓"):
    """A vertical connector between diagram boxes."""
    p = _para(doc, space_before=1, space_after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, glyph, size=SZ_BODY, font=ARROW_FONT)
    if label:
        _run(p, "  " + label, size=SZ_TABLE)
    return p


def dsplit(doc, label=""):
    """A one-to-many connector."""
    p = _para(doc, space_before=1, space_after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "↓      ↓      ↓      ↓", size=SZ_BODY, font=ARROW_FONT)
    if label:
        _run(p, "  " + label, size=SZ_TABLE)
    return p


# ---------------------------------------------------------------------------
# Page furniture
# ---------------------------------------------------------------------------

def setup_page(doc):
    for s in doc.sections:
        s.page_width = Inches(8.5)
        s.page_height = Inches(11.0)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.header_distance = Inches(0.49)
        s.footer_distance = Inches(0.49)
        s.different_first_page_header_footer = True


def add_page_number_footer(doc):
    """Footer reading 'หน้า N'. The title page is left blank via w:titlePg."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "หน้า ", size=SZ_SMALL)

    r = p.add_run()
    r.font.name = THAI_FONT
    r.font.size = Pt(SZ_SMALL)
    r.font.color.rgb = BLACK
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), THAI_FONT)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(SZ_SMALL * 2)))
    rPr.append(szCs)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._element.append(begin)
    r._element.append(instr)
    r._element.append(end)


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------

def cover(doc):
    p = _para(doc, space_before=82, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "SMARTLIVA", size=SZ_SUBTITLE, bold=True)

    p = _para(doc, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER,
              line_spacing=1.2)
    _run(p, "เอกสารเปิดเผยรายละเอียดทางเทคนิค", size=SZ_TITLE, bold=True)

    p = _para(doc, space_before=0, space_after=14, align=WD_ALIGN_PARAGRAPH.CENTER,
              line_spacing=1.2)
    _run(p, "ของโมเดลการเรียนรู้ที่ออกแบบขึ้นใหม่", size=SZ_TITLE, bold=True)

    p = _para(doc, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "โครงการระบบปัญญาประดิษฐ์ช่วยคัดกรองภาพอัลตราซาวด์ตับชนิด B-mode",
         size=SZ_SUBTITLE)

    p = _para(doc, space_before=0, space_after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "จัดทำเพื่อประกอบการพิจารณาแนวทางการขอรับความคุ้มครองทรัพย์สินทางปัญญาเพิ่มเติม",
         size=SZ_SUBTITLE)

    rule(doc, space_before=0, space_after=14)

    meta = [
        ("ขอบเขตเนื้อหา", "ภาค ก พังผืดตับและไขมันพอกตับ · ภาค ข รอยโรค มะเร็ง และพยาธิใบไม้ตับ"),
        ("ผลงานเดิมที่ใช้เทียบ", "SmartLiva-LiverUS-SDK เวอร์ชัน 1.1.0"),
        ("ฐานอ้างอิง", "ซอร์สโค้ด ไฟล์น้ำหนักโมเดล และบันทึกการฝึกที่ตรวจสอบได้ในโครงการ"),
        ("สถานะเอกสาร", "ร่างเพื่อการพิจารณาภายใน · ยังไม่เผยแพร่สู่สาธารณะ"),
        ("วันที่จัดทำ", "27 สิงหาคม 2569"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for r, (k, v) in enumerate(meta):
        ck, cv = t.cell(r, 0), t.cell(r, 1)
        ck.width, cv.width = Inches(1.85), Inches(4.65)
        _set_cell_margins(ck, top=60, bottom=60, left=0, right=110)
        _set_cell_margins(cv, top=60, bottom=60, left=0, right=0)
        _cell_text(ck, k, size=SZ_SMALL, bold=True)
        _cell_text(cv, v, size=SZ_SMALL)

    p = _para(doc, space_before=26, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "เอกสารขาว-ดำ | แบบอักษร TH SarabunPSK", size=SZ_SMALL)

    page_break(doc)


def contents(doc):
    h1(doc, "สารบัญ")
    entries = [
        ("1.", "วัตถุประสงค์ ขอบเขต และวิธีจัดทำเอกสาร", 0),
        ("2.", "ภาพรวมระบบและส่วนที่ต่างไปจากผลงานเดิม", 0),
        ("3.", "ลักษณะการทำงานของระบบในการรับเข้าข้อมูล", 0),
        ("", "ภาค ก — โมเดลพังผืดตับและไขมันพอกตับ", 1),
        ("4.", "โมเดลจัดระดับพังผืดตับ F0–F4 (FibrosisNet Ensemble)", 0),
        ("5.", "โมดูลประเมินไขมันพอกตับ S0–S3", 0),
        ("6.", "จุดร่วมภายในภาค ก", 0),
        ("", "ภาค ข — โมเดลรอยโรค มะเร็ง และพยาธิใบไม้ตับ", 1),
        ("7.", "ตัวตรวจจับรอยโรคเฉพาะที่ในเนื้อตับ", 0),
        ("8.", "กฎยกระดับเมื่อพบรอยโรคมะเร็ง และการส่งตรวจยืนยัน", 0),
        ("9.", "โมดูลประเมินความเสี่ยงพยาธิใบไม้ตับและมะเร็งท่อน้ำดี", 0),
        ("10.", "จุดร่วมภายในภาค ข", 0),
        ("", "ส่วนที่ 3 — จุดร่วมของแต่ละโมเดล", 1),
        ("11.", "จุดร่วมของทั้งสี่โมดูล", 0),
        ("", "ภาคผนวก", 1),
        ("ก.", "ตารางค่าคงที่และเกณฑ์ตัดสินทั้งหมด", 0),
        ("ข.", "ดัชนีไฟล์และตำแหน่งบรรทัดอ้างอิง", 0),
        ("ค.", "ประเด็นที่ต้องยืนยันก่อนยกร่างข้อถือสิทธิ", 0),
    ]
    for num, text, kind in entries:
        if kind == 1:
            p = _para(doc, space_before=8, space_after=3, indent=0.0)
            _run(p, text, size=SZ_BODY, bold=True)
        else:
            p = _para(doc, space_after=3, indent=0.42)
            p.paragraph_format.first_line_indent = Inches(-0.42)
            _run(p, f"{num}\t", size=SZ_BODY, bold=True)
            _run(p, text, size=SZ_BODY)
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 1
# ---------------------------------------------------------------------------

def section_1(doc):
    h1(doc, "1. วัตถุประสงค์ ขอบเขต และวิธีจัดทำเอกสาร")

    h2(doc, "1.1 คำถามที่เอกสารนี้ตอบ")
    body(doc, "เอกสารฉบับนี้จัดทำขึ้นเพื่อตอบคำถามสามข้อที่ได้รับจากฝ่ายทรัพย์สินทางปัญญา "
              "สำหรับใช้ประเมินความเป็นไปได้และแนวทางการขอรับความคุ้มครองเพิ่มเติม ดังนี้")
    numbered(doc, 1, "โครงสร้าง ลักษณะการทำงาน และลักษณะทางเทคนิคของโมเดลการเรียนรู้ที่ออกแบบขึ้นใหม่ "
                     "สำหรับการตรวจหารอยโรค พังผืด มะเร็ง หรือพยาธิ — ตอบไว้ในภาค ก (หัวข้อ 4–6) "
                     "และภาค ข (หัวข้อ 7–10)")
    numbered(doc, 2, "ลักษณะการทำงานของระบบในการรับเข้าข้อมูลเพื่อประมวลผลและประเมินความเสี่ยง "
                     "หรือการระบุพื้นที่ของรอยโรค เฉพาะส่วนที่ปรับเปลี่ยนไปจากข้อมูลเดิม — "
                     "ตอบไว้ในหัวข้อ 2 และหัวข้อ 3")
    numbered(doc, 3, "จุดร่วมของแต่ละโมเดล — ตอบไว้ในหัวข้อ 6 หัวข้อ 10 และหัวข้อ 11")

    h2(doc, "1.2 นิยามของ “ข้อมูลเดิม” ที่ใช้เป็นฐานเปรียบเทียบ")
    body(doc, "คำว่า “ข้อมูลเดิม” ในเอกสารนี้หมายถึงแพ็กเกจ SmartLiva-LiverUS-SDK เวอร์ชัน 1.1.0 "
              "ซึ่งเป็นผลงานที่มีอยู่ก่อนและได้เปิดเผยรายละเอียดไว้แล้ว แพ็กเกจดังกล่าวทำหน้าที่เพียง "
              "สามอย่างคือ ตรวจว่าภาพที่รับเข้าเป็นภาพอัลตราซาวด์ B-mode จริงหรือไม่ จำแนกว่าเป็นอวัยวะใด "
              "ในสิบชนิด และวาดเส้นขอบเขตตับออกมาเป็นพิกัดรูปหลายเหลี่ยม "
              "โดยคู่มือของแพ็กเกจระบุไว้ชัดเจนว่าการบอกว่าตับปกติหรือเป็นโรค ไม่ว่าจะเป็นไขมัน พังผืด "
              "ก้อนเนื้อ หรือมะเร็ง ไม่ใช่ขอบเขตงานของแพ็กเกจนั้น")
    body(doc, "ด้วยเหตุนี้ โมดูลวินิจฉัยโรคทุกตัวที่บรรยายในเอกสารฉบับนี้จึงเป็นส่วนที่สร้างขึ้นใหม่ทั้งหมด "
              "ไม่ได้ดัดแปลงจากส่วนที่เปิดเผยไว้เดิม และเป็นสาระสำคัญของการขอรับความคุ้มครองเพิ่มเติม "
              "ส่วนที่ต่อยอดจากของเดิมมีเพียงชั้นคัดกรองภาพและการสร้างขอบเขตตับ ซึ่งเอกสารนี้อธิบายไว้ใน "
              "หัวข้อ 2 และหัวข้อ 3 พร้อมระบุจุดที่ปรับเปลี่ยนไป")

    h2(doc, "1.3 วิธีจัดทำและระดับความน่าเชื่อถือของตัวเลข")
    body(doc, "ทุกค่าคงที่ สูตรคำนวณ เกณฑ์ตัดสิน และตัวเลขประสิทธิภาพในเอกสารนี้ อ่านออกมาจากซอร์สโค้ด "
              "ไฟล์น้ำหนักโมเดล ไฟล์ตั้งค่าการฝึก และบันทึกผลการวัดที่มีอยู่จริงในโครงการ ณ วันที่จัดทำ "
              "ภาคผนวก ข ระบุตำแหน่งไฟล์และบรรทัดของตรรกะสำคัญทุกจุดไว้ให้ตรวจสอบย้อนกลับได้")
    body(doc, "เอกสารนี้ยึดหลักเดียวกับรายงานแหล่งที่มาชุดข้อมูลที่ส่งไปก่อนหน้า คือแยกให้ชัดว่าโมดูลใด "
              "เป็นโมเดลที่ฝึกด้วยข้อมูล และโมดูลใดเป็นกฎคำนวณเชิงฟิสิกส์หรือกฎจากประวัติผู้ป่วย "
              "เนื่องจากลักษณะทางเทคนิคและแนวทางการขอความคุ้มครองของทั้งสองแบบไม่เหมือนกัน")
    lead(doc, "ข้อควรทราบ: ",
         "มีข้อความบางส่วนในเอกสารประกอบเดิมของโครงการที่บรรยายความสามารถซึ่งยังไม่ปรากฏใน "
         "ซอร์สโค้ดที่ใช้งานจริง เอกสารฉบับนี้ไม่นำข้อความเหล่านั้นมาใช้ และได้รวบรวมไว้ในภาคผนวก ค "
         "เพื่อให้ผู้ยกร่างคำขอเห็นก่อนกำหนดขอบเขตข้อถือสิทธิ")

    h2(doc, "1.4 คำย่อที่ใช้ในเอกสาร")
    table(
        doc,
        ["คำย่อ", "ความหมาย"],
        [
            ("B-mode", "ภาพอัลตราซาวด์แบบสร้างความสว่างตามความเข้มสัญญาณสะท้อน"),
            ("F0–F4", "ระดับพังผืดในตับตามระบบ METAVIR ตั้งแต่ไม่มีพังผืดถึงตับแข็ง"),
            ("S0–S3", "ระดับไขมันพอกตับ ตั้งแต่ไม่พบถึงระดับรุนแรง"),
            ("kPa", "หน่วยกิโลพาสคาล ใช้บอกความแข็งของเนื้อตับ"),
            ("TE", "การวัดความแข็งตับด้วยคลื่นยืดหยุ่นชนิด Transient Elastography"),
            ("SWE", "การวัดความแข็งตับด้วยคลื่นเฉือน Shear Wave Elastography"),
            ("ROI", "บริเวณที่สนใจซึ่งถูกตัดออกมาเพื่อวิเคราะห์"),
            ("mask", "ภาพหน้ากากสองระดับที่ระบุว่าพิกเซลใดเป็นเนื้อตับ"),
            ("CORN", "วิธีจำแนกแบบลำดับขั้นชนิด Conditional Ordinal Regression for Neural networks"),
            ("NMS", "การกดทับกล่องตรวจจับที่ซ้อนทับกัน Non-Maximum Suppression"),
            ("HCC", "มะเร็งตับชนิดปฐมภูมิ Hepatocellular Carcinoma"),
            ("CCA", "มะเร็งท่อน้ำดี Cholangiocarcinoma"),
            ("FFC / FFS", "การเปลี่ยนแปลงไขมันเฉพาะที่ และการเว้นไขมันเฉพาะที่"),
        ],
        widths=[1.35, 5.15],
    )
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 2
# ---------------------------------------------------------------------------

def section_2(doc):
    h1(doc, "2. ภาพรวมระบบและส่วนที่ต่างไปจากผลงานเดิม")

    h2(doc, "2.1 ขอบเขตของผลงานเดิม (SmartLiva-LiverUS-SDK v1.1.0)")
    body(doc, "ผลงานเดิมเป็นแพ็กเกจคัดกรองภาพที่ทำงานเป็นทอดสองขั้น ขั้นแรกเป็นด่านหน้าที่ตัดสินว่าภาพ "
              "ที่รับเข้าใช้วิเคราะห์ได้หรือไม่ และเป็นอวัยวะใด ขั้นที่สองคือการวาดเส้นขอบตับเมื่อผ่านด่านแรกแล้ว "
              "องค์ประกอบทั้งหมดมีดังนี้")
    bullet(doc, "ตรวจการกระจายตัวของจุดรบกวนแบบ speckle เทียบกับกรอบค่าที่ปรับเทียบไว้ "
                "แยกตามรูปทรงหัวตรวจชนิดเชิงเส้นและชนิดพัด", label="ด่านฟิสิกส์ของภาพ: ")
    bullet(doc, "โครงข่าย ResNet-18 จำแนกอวัยวะสิบชนิด ได้แก่ เต้านม หลอดเลือดคาโรติด ทารกในครรภ์ "
                "ถุงน้ำดี ไต ตับ อื่น ๆ ตับอ่อน ม้าม และไทรอยด์ ความถูกต้อง 0.9932 บนชุดทดสอบข้ามแหล่ง "
                "จำนวน 592 ภาพ", label="ด่านจำแนกอวัยวะ: ")
    bullet(doc, "โครงข่าย U-Net สามคลาส คือ พื้นหลัง ตับ และถุงน้ำดี ทำนายที่ความละเอียด 256×256 "
                "แล้วขยายกลับเป็นขนาดภาพจริง ค่า Dice ของตับเท่ากับ 0.9333 ในการทดสอบภายในแหล่งเดียวกัน "
                "และ 0.891 เมื่อทดสอบข้ามโรงพยาบาล ส่วนถุงน้ำดีได้ 0.8785",
                label="การวาดขอบเขต: ")
    bullet(doc, "ระบุว่าเป็นภาพตับหรือไม่ ค่าความมั่นใจ สถานะคุณภาพ และพิกัดรูปหลายเหลี่ยม "
                "ของขอบตับในระบบพิกัดของภาพต้นฉบับ", label="ผลลัพธ์ที่ส่งออก: ")
    body(doc, "คู่มือของแพ็กเกจระบุไว้เป็นข้อความชัดเจนว่าแพ็กเกจนี้ไม่ตอบว่าตับปกติหรือเป็นโรค และเป็น "
              "ด่านคัดกรองภาพ ไม่ใช่เครื่องมือวินิจฉัย ขอบเขตนี้จึงเป็นเส้นแบ่งที่ชัดเจนระหว่างผลงานเดิม "
              "กับส่วนที่ขอความคุ้มครองเพิ่มในเอกสารฉบับนี้")

    h2(doc, "2.2 ขอบเขตของระบบใหม่")
    body(doc, "ระบบใหม่รับส่วนคัดกรองภาพและการวาดขอบตับของเดิมมาเป็นชั้นล่างสุด แล้วสร้างชั้นวินิจฉัยโรค "
              "ขึ้นใหม่ทั้งหมดบนขอบเขตตับนั้น ประกอบด้วยโมดูลผู้เชี่ยวชาญสี่ตัวที่ทำงานบนหน้ากากตับชุดเดียวกัน "
              "ตามด้วยชั้นรวมผลเชิงกฎและชั้นเรียบเรียงรายงาน")
    table(
        doc,
        ["โมดูลใหม่", "ประเภททางเทคนิค", "ผลลัพธ์"],
        [
            ("พังผืดตับ", "โครงข่ายประสาทเทียมที่ฝึกด้วยข้อมูล\n(ensemble 5 โมเดล)",
             "ระดับ F0–F4, ค่าประมาณ kPa,\nความน่าจะเป็นสะสม, ระดับความเสี่ยง"),
            ("ไขมันพอกตับ", "การคำนวณเชิงฟิสิกส์แบบกำหนดได้\n(ไม่มีไฟล์น้ำหนัก)",
             "ระดับ S0–S3, อัตราการลดทอนลำคลื่น"),
            ("รอยโรคเฉพาะที่", "โครงข่ายตรวจจับวัตถุที่ฝึกด้วยข้อมูล\n(YOLOv8n, 7 คลาส)",
             "กล่องขอบเขต ชนิดรอยโรค ค่าความมั่นใจ\nและขนาดโดยประมาณ"),
            ("ความเสี่ยงพยาธิใบไม้ตับ", "กฎให้คะแนนเชิงเส้นจากประวัติผู้ป่วย\n(ไม่ได้เรียนรู้จากภาพ)",
             "Negative / Possible / Probable\nและระดับความเสี่ยง"),
        ],
        widths=[1.55, 2.45, 2.5],
    )

    h2(doc, "2.3 ตารางเปรียบเทียบส่วนที่ต่างไปจากข้อมูลเดิม")
    table(
        doc,
        ["หัวข้อ", "ผลงานเดิม v1.1.0", "ระบบใหม่"],
        [
            ("คำถามที่ตอบได้", "เป็นภาพ US จริงหรือไม่ · อวัยวะใด ·\nขอบตับอยู่ตรงไหน",
             "เพิ่มการตอบว่าเป็นโรคใด ระดับใด\nและรอยโรคอยู่ตำแหน่งใด"),
            ("ข้อมูลที่รับเข้า", "ไฟล์ภาพอย่างเดียว",
             "ไฟล์ภาพ พร้อมข้อมูลห้องปฏิบัติการ\nค่าความแข็งตับที่วัดมา และประวัติผู้ป่วย"),
            ("การลบข้อมูลระบุตัวตน", "ไม่มีในสายงาน",
             "ลบแถบบนของเฟรม 12 เปอร์เซ็นต์\nก่อนเข้าสู่การวิเคราะห์"),
            ("การสร้างขอบเขตตับ", "U-Net สามคลาสอย่างเดียว",
             "U-Net ทำหน้าที่สร้างคำใบ้และคัดกรอง\nร่วมกับ MedSAM2 ที่สร้างหน้ากาก"),
            ("การหักถุงน้ำดี", "ทำนายไว้ภายในแต่ไม่รายงานออก",
             "หักพิกเซลถุงน้ำดีออกจากเนื้อตับสองรอบ\nเพื่อให้ทุกโมดูลใช้เนื้อตับล้วน"),
            ("ด่านกันการทำงานผิดบริบท", "หยุดเมื่อไม่ใช่ตับ",
             "หยุดเมื่อไม่ใช่ตับ และเพิ่มเงื่อนไข\nสัดส่วนพื้นที่ตับน้อยกว่า 5 เปอร์เซ็นต์"),
            ("การระบุพื้นที่รอยโรค", "ไม่มี",
             "กล่องขอบเขตที่ผ่านการกักบริเวณด้วยหน้ากากตับ\nรายงานเป็นพิกัดปรับมาตรฐาน 0 ถึง 1"),
            ("การรวมผลข้ามโมดูล", "ไม่มี",
             "ชั้นกฎรวมผล ยกระดับความเสี่ยง\nและออกคำแนะนำส่งตรวจยืนยัน"),
            ("วงจรแพทย์ตรวจทาน", "ไม่มี",
             "บันทึกการยืนยันและการแก้ไขของแพทย์\nแบบต่อท้ายอย่างเดียว เพื่อใช้ฝึกซ้ำ"),
        ],
        widths=[1.45, 2.45, 2.6],
    )
    page_break(doc)

    h2(doc, "2.4 แผนภาพที่ 1 เปรียบเทียบขอบเขตของเดิมกับระบบใหม่")
    drow(
        doc,
        [
            ["ผลงานเดิม v1.1.0",
             "ภาพอัลตราซาวด์",
             "ด่านฟิสิกส์ของภาพ",
             "ด่านจำแนกอวัยวะ 10 คลาส",
             "U-Net วาดขอบตับ 3 คลาส",
             "พิกัดรูปหลายเหลี่ยมของขอบตับ",
             "— จบเพียงเท่านี้ —"],
            ["ระบบใหม่",
             "ภาพอัลตราซาวด์ + ข้อมูลคลินิก",
             "ลบข้อมูลระบุตัวตน",
             "ด่านฟิสิกส์ + ด่านอวัยวะ + เกณฑ์พื้นที่ตับ",
             "U-Net ร่วมกับ MedSAM2 และหักถุงน้ำดี",
             "โมดูลวินิจฉัยโรค 4 ตัวบนหน้ากากตับชุดเดียว",
             "ชั้นกฎรวมผล และรายงานพร้อมคำแนะนำ"],
        ],
        widths=[3.0, 3.5],
    )
    caption(doc, "แผนภาพที่ 1 กล่องคอลัมน์ซ้ายคือขอบเขตที่เปิดเผยไว้เดิม ส่วนคอลัมน์ขวาคือระบบใหม่ "
                 "โดยสามบรรทัดล่างของคอลัมน์ขวาไม่มีคู่เทียบในผลงานเดิม")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 3
# ---------------------------------------------------------------------------

def section_3(doc):
    h1(doc, "3. ลักษณะการทำงานของระบบในการรับเข้าข้อมูล")
    body(doc, "หัวข้อนี้ตอบคำถามข้อ 2 โดยตรง คือ อธิบายว่าระบบรับข้อมูลเข้ามาอย่างไร แปลงข้อมูลนั้นเป็นอะไร "
              "ก่อนส่งต่อให้โมดูลวินิจฉัย และระบุพื้นที่ของรอยโรคออกมาในรูปแบบใด พร้อมชี้จุดที่ต่างไปจากเดิม")

    h2(doc, "3.1 ช่องทางรับเข้าข้อมูล")
    body(doc, "ระบบเปิดช่องทางรับเข้าสองแบบ แบบแรกคือช่องทางวิเคราะห์เต็มรูปแบบซึ่งเรียกใช้โมดูลทั้งสี่ "
              "พร้อมชั้นรวมผล และแบบที่สองคือช่องทางรายโมดูลสำหรับเรียกใช้เฉพาะโมดูลใดโมดูลหนึ่ง")
    table(
        doc,
        ["ช่องทาง", "ข้อมูลที่รับเข้า", "ขอบเขตการทำงาน"],
        [
            ("วิเคราะห์เต็มรูปแบบ", "ไฟล์ภาพ และข้อมูลคลินิกในรูปแบบ JSON",
             "ด่านคัดกรอง → สร้างหน้ากากตับ →\nโมดูลทั้งสี่ → ชั้นกฎรวมผล → รายงาน"),
            ("รายโมดูล พังผืด", "ไฟล์ภาพ และมุมตรวจ (ไม่บังคับ)", "ด่านอวัยวะ → หน้ากากตับ → โมดูลพังผืด"),
            ("รายโมดูล รอยโรค", "ไฟล์ภาพ และค่าความมั่นใจขั้นต่ำ", "ด่านอวัยวะ → หน้ากากตับ → โมดูลรอยโรค"),
            ("รายโมดูล ไขมัน", "ไฟล์ภาพ", "ด่านอวัยวะ → หน้ากากตับ → โมดูลไขมัน"),
            ("รายโมดูล พยาธิ", "ไฟล์ภาพ และประวัติผู้ป่วยในรูปแบบ JSON", "ด่านอวัยวะ → หน้ากากตับ → โมดูลพยาธิ"),
            ("คัดกรองและวาดขอบ", "ไฟล์ภาพ",
             "เทียบเท่าขอบเขตของผลงานเดิม\nไม่เรียกโมดูลวินิจฉัยใด ๆ"),
        ],
        widths=[1.6, 2.35, 2.55],
    )
    lead(doc, "ข้อแตกต่างจากเดิม: ",
         "ผลงานเดิมรับเข้าเฉพาะไฟล์ภาพ ส่วนระบบใหม่รับข้อมูลทางคลินิกควบคู่ไปกับภาพ และข้อมูลนั้น "
         "มีผลต่อผลลัพธ์จริง เช่น มุมตรวจถูกป้อนเข้าโมเดลพังผืดโดยตรง และประวัติการกินปลาน้ำจืดดิบ "
         "เป็นตัวแปรหลักของโมดูลความเสี่ยงพยาธิใบไม้ตับ")

    h2(doc, "3.2 รูปแบบภาพที่รองรับและข้อจำกัดที่ตามมา")
    bullet(doc, "PNG, JPEG, WebP และ BMP โดยแปลงเป็นภาพสีสามช่องสัญญาณก่อนประมวลผลทุกกรณี",
           label="รูปแบบที่รองรับ: ")
    bullet(doc, "25 เมกะไบต์ต่อไฟล์ หากเกินจะถูกปฏิเสธก่อนถอดรหัสภาพ", label="ขนาดสูงสุด: ")
    bullet(doc, "ระบบไม่รองรับไฟล์ DICOM จึงไม่มีค่าระยะห่างต่อพิกเซลจากเครื่องตรวจ "
                "ผลที่ตามมาคือการรายงานขนาดรอยโรคต้องใช้ค่าคงที่แทน ดังอธิบายในหัวข้อ 7.5 "
                "และระบบไม่รายงานพื้นที่เป็นตารางเซนติเมตร", label="ข้อจำกัดสำคัญ: ")

    h2(doc, "3.3 ข้อมูลทางคลินิกที่รับเข้าและการใช้งาน")
    table(
        doc,
        ["กลุ่มข้อมูล", "รายการ", "ถูกใช้โดย"],
        [
            ("ผลห้องปฏิบัติการ", "AST, ALT, เกล็ดเลือด, บิลิรูบิน",
             "ชั้นตรวจความสมเหตุสมผลของข้อมูล\n(ตรวจค่าติดลบ)"),
            ("ค่าความแข็งตับที่วัดมา", "ค่า kPa จาก TE และค่า CAP",
             "แปลงเป็นระดับ F เพื่อแสดงเทียบเคียง\nและตรวจค่าที่เกิน 75 kPa ว่าผิดปกติ"),
            ("ประวัติผู้ป่วย", "อายุ, เพศ, ไวรัสตับอักเสบบี, ไวรัสตับอักเสบซี,\nประวัติดื่มสุรา, การกินปลาน้ำจืดดิบ",
             "การกินปลาน้ำจืดดิบเป็นตัวแปรหลัก\nของโมดูลความเสี่ยงพยาธิใบไม้ตับ"),
            ("พารามิเตอร์การวิเคราะห์", "มุมตรวจ และค่าความมั่นใจขั้นต่ำของตัวตรวจจับ",
             "มุมตรวจป้อนเข้าโมเดลพังผืด\nค่าความมั่นใจใช้กรองผลตัวตรวจจับ"),
        ],
        widths=[1.5, 2.75, 2.25],
    )
    body(doc, "มุมตรวจที่ระบบรู้จักมีเจ็ดค่า ตรงกับมุมมาตรฐานที่ใช้เก็บภาพในโครงการ ได้แก่ RH, GBH, LHA, "
              "LHP, SPH, LHV และ FPH หากผู้ใช้ไม่ระบุ ระบบจะถือว่าเป็นมุมที่ไม่ทราบ ซึ่งโมเดลพังผืดรองรับ "
              "ด้วยตำแหน่งเวกเตอร์ฝังเฉพาะสำหรับกรณีนี้")
    page_break(doc)

    h2(doc, "3.4 การลบข้อมูลระบุตัวตนก่อนวิเคราะห์")
    body(doc, "ทันทีที่ถอดรหัสภาพเสร็จ ระบบจะกำหนดค่าพิกเซลของแถวบนสุดจำนวน 12 เปอร์เซ็นต์ของความสูงเฟรม "
              "ให้เป็นศูนย์ เพื่อลบแถบข้อความที่เครื่องอัลตราซาวด์พิมพ์ทับไว้ ซึ่งโดยทั่วไปบรรจุชื่อผู้ป่วย "
              "วันที่ตรวจ และรหัสสถานพยาบาล ภาพระดับสีเทาและภาพสีที่ส่งต่อให้ทุกโมดูลวินิจฉัยล้วนเป็นภาพ "
              "ที่ผ่านการลบแถบนี้แล้ว")
    body(doc, "ขั้นตอนนี้ไม่มีในผลงานเดิม และมีผลข้างเคียงเชิงเทคนิคที่ตั้งใจ คือช่วยลดโอกาสที่โมเดลจะเรียนรู้ "
              "หรืออนุมานจากตัวอักษรที่เครื่องพิมพ์ทับ แทนที่จะเรียนรู้จากลักษณะของเนื้อตับจริง")

    h2(doc, "3.5 ด่านบังคับก่อนเข้าสู่การประมวลผลโรค")
    body(doc, "ระบบกำหนดให้ทุกภาพต้องผ่านด่านสองชั้นก่อน จึงจะมีโมดูลวินิจฉัยใดทำงานได้ "
              "หากไม่ผ่านด่านใดด่านหนึ่ง ระบบจะคืนผลลัพธ์ที่ระบุสถานะหยุดทำงาน "
              "โดยที่ยังไม่มีการสร้างงานประมวลผลของโมดูลใดขึ้นเลย ไม่ใช่การสร้างขึ้นแล้วทิ้งผลภายหลัง")
    numbered(doc, 1, "ต้องจำแนกได้ว่าเป็นตับ โดยค่าความมั่นใจต้องไม่ต่ำกว่า 0.55 และค่าเอนโทรปีของ "
                     "การแจกแจงความน่าจะเป็นต้องไม่เกิน 1.30 หากไม่เข้าเงื่อนไข ระบบจะตัดสินว่าไม่แน่ใจ "
                     "และไม่ถือว่าเป็นตับ", label="ด่านจำแนกอวัยวะ: ")
    numbered(doc, 2, "หลังสร้างหน้ากากตับแล้ว สัดส่วนพิกเซลตับต่อพื้นที่เฟรมทั้งหมดต้องไม่น้อยกว่า "
                     "5.0 เปอร์เซ็นต์ มิฉะนั้นถือว่าภาพมีเนื้อตับไม่พอสำหรับการวินิจฉัย",
                     label="ด่านสัดส่วนพื้นที่ตับ: ")
    lead(doc, "เหตุผลเชิงเทคนิค: ",
         "โมเดลวินิจฉัยทุกตัวถูกฝึกด้วยภาพตับล้วน หากป้อนภาพอวัยวะอื่นเข้าไปโดยตรง โมเดลจะให้คำตอบ "
         "ที่ดูสมเหตุสมผลแต่ไม่มีความหมาย ด่านนี้จึงเป็นกลไกกันความผิดพลาดเชิงบริบทที่บังคับใช้ในระดับ "
         "โครงสร้างของสายงาน ไม่ใช่คำเตือนที่ผู้ใช้เลือกมองข้ามได้")

    h2(doc, "3.6 การสร้างขอบเขตตับที่ใช้ร่วมกัน")
    body(doc, "ระบบใหม่เปลี่ยนวิธีสร้างหน้ากากตับจากเดิมที่ใช้ U-Net เพียงตัวเดียว มาเป็นการทำงานร่วมกัน "
              "ของสองโมเดล โดยแบ่งบทบาทกันชัดเจน คือ U-Net ทำหน้าที่สร้างคำใบ้และคัดกรอง "
              "ส่วน MedSAM2 ทำหน้าที่สร้างหน้ากากจริง ลำดับการทำงานมีดังนี้")
    numbered(doc, 1, "สกัดกรวยสัญญาณอัลตราซาวด์จากเฟรม เพื่อกำหนดขอบเขตพื้นที่ที่มีสัญญาณจริง")
    numbered(doc, 2, "ให้ U-Net สามคลาสทำนายเมล็ดเริ่มต้นของตับและถุงน้ำดี โดยรับเมล็ดตับเมื่อมีพิกเซล "
                     "เกิน 100 จุด และรับเมล็ดถุงน้ำดีเมื่อเกิน 60 จุด")
    numbered(doc, 3, "แปลงจุดศูนย์กลางของถุงน้ำดีเป็นจุดชี้เชิงลบ เพื่อบอก MedSAM2 ว่าบริเวณนั้นไม่ใช่ตับ")
    numbered(doc, 4, "เลือกกล่องคำใบ้ตามลำดับความสำคัญ คือ ใช้กล่องจากตัวตรวจจับตับก่อน "
                     "หากไม่มีจึงใช้กรอบของเมล็ดจาก U-Net ถัดไปใช้กรอบของกรวยสัญญาณ "
                     "และสุดท้ายใช้กรอบคงที่ที่ครอบร้อยละ 15 ถึง 85 ของเฟรม")
    numbered(doc, 5, "เรียก MedSAM2 ด้วยกล่องคำใบ้หนึ่งกล่อง จุดชี้เชิงบวกหนึ่งจุดที่จุดกึ่งกลางกล่อง "
                     "และจุดชี้เชิงลบตามจำนวนถุงน้ำดีที่พบ แล้วเลือกหน้ากากที่ได้คะแนนสูงสุด")
    numbered(doc, 6, "หักพิกเซลถุงน้ำดีออกจากหน้ากากที่ได้ จำกัดผลให้อยู่ภายในกรวยสัญญาณ "
                     "และตัดพิกเซลที่มืดกว่าระดับ 8 ออกเพื่อกันเงาอะคูสติก")
    numbered(doc, 7, "ปรับแต่งทางสัณฐานวิทยา คือปิดช่องว่าง เก็บเฉพาะองค์ประกอบเชื่อมต่อที่ใหญ่ที่สุด "
                     "และเติมรูภายใน จากนั้นหักพิกเซลถุงน้ำดีที่ผ่านการปรับแต่งแล้วออกอีกรอบหนึ่ง")
    lead(doc, "ผลลัพธ์เชิงคุณสมบัติ: ",
         "หน้ากากตับที่ได้ไม่มีพิกเซลถุงน้ำดีปนอยู่เลยโดยโครงสร้าง เนื่องจากมีการหักออกสองรอบ "
         "คุณสมบัตินี้สำคัญกับทุกโมดูลถัดไป เพราะถุงน้ำดีมีความสว่างต่างจากเนื้อตับมาก "
         "หากปนอยู่จะทำให้ค่าสถิติความสว่างของโมดูลไขมันเพี้ยน และทำให้ตัวตรวจจับรอยโรค "
         "รายงานสิ่งที่อยู่ในถุงน้ำดีว่าเป็นรอยโรคในเนื้อตับ")

    h2(doc, "3.7 รูปแบบการระบุพื้นที่ที่ส่งออก")
    body(doc, "ทุกโมดูลรายงานตำแหน่งด้วยโครงสร้างข้อมูลชุดเดียวกัน โดยพิกัดทุกค่าถูกปรับมาตรฐานให้อยู่ในช่วง "
              "0 ถึง 1 เทียบกับความกว้างและความสูงของเฟรม จึงไม่ผูกกับความละเอียดของภาพต้นฉบับ "
              "และนำไปวางซ้อนบนภาพขนาดใดก็ได้")
    table(
        doc,
        ["เขตข้อมูล", "ความหมาย"],
        [
            ("regionId", "รหัสอ้างอิงพื้นที่ ใช้คำนำหน้าคงที่ตามโมดูลที่สร้าง"),
            ("shape", "รูปทรง เช่น กล่องสองจุด หรือเส้นอิสระตั้งแต่สองจุดขึ้นไป"),
            ("points", "พิกัดปรับมาตรฐาน 0 ถึง 1 กล่องใช้สองคู่คือมุมบนซ้ายและมุมล่างขวา"),
            ("label", "ข้อความกำกับที่แสดงต่อผู้ใช้"),
            ("confidence", "ค่าความมั่นใจของพื้นที่นั้น"),
            ("source", "โมดูลต้นทาง ใช้แยกว่าเป็นผลของ AI หรือของแพทย์ที่แก้ไขเข้ามา"),
        ],
        widths=[1.5, 5.0],
    )
    page_break(doc)

    h2(doc, "3.8 แผนภาพที่ 2 ผังการรับเข้าข้อมูลและการไหลของข้อมูล")
    dbox(doc, ["ข้อมูลเข้า", "ไฟล์ภาพอัลตราซาวด์ B-mode  +  ข้อมูลคลินิกรูปแบบ JSON"], width=5.6)
    darrow(doc)
    dbox(doc, ["ขั้นเตรียมข้อมูล", "ถอดรหัสภาพ · ตรวจขนาดไม่เกิน 25 MB · ลบแถบบน 12 เปอร์เซ็นต์",
               "แปลงเป็นภาพระดับสีเทาและภาพสีสำหรับแต่ละโมดูล"], width=5.6)
    darrow(doc)
    dbox(doc, ["ด่านคัดกรอง (บังคับ)",
               "ด่านคุณภาพภาพ → ด่านจำแนกอวัยวะ 10 คลาส",
               "ไม่ผ่าน → หยุดทันที ไม่มีโมดูลวินิจฉัยใดทำงาน"], width=5.6, heavy=True)
    darrow(doc)
    dbox(doc, ["สร้างขอบเขตตับชุดเดียว",
               "U-Net สร้างคำใบ้และจุดชี้เชิงลบ → MedSAM2 สร้างหน้ากาก",
               "หักพิกเซลถุงน้ำดีออกสองรอบ · ปรับแต่งทางสัณฐานวิทยา"], width=5.6)
    darrow(doc, "ตรวจสัดส่วนพื้นที่ตับ ต้องไม่น้อยกว่า 5.0 เปอร์เซ็นต์")
    drow(
        doc,
        [
            ["พังผืด", "F0–F4", "และค่า kPa"],
            ["ไขมัน", "S0–S3", "จากค่าฟิสิกส์"],
            ["รอยโรค", "กล่องขอบเขต", "7 ชนิด"],
            ["พยาธิ", "ระดับความเสี่ยง", "จากประวัติ"],
        ],
        widths=[1.6, 1.6, 1.65, 1.65],
    )
    darrow(doc)
    dbox(doc, ["ชั้นกฎรวมผลเชิงกำหนด",
               "รวมคำเตือน · ยกระดับความเสี่ยงเมื่อพบรอยโรคมะเร็ง",
               "ออกคำแนะนำส่งตรวจยืนยัน"], width=5.6)
    darrow(doc)
    dbox(doc, ["ผลลัพธ์", "รายงานข้อความ · พื้นที่พิกัดปรับมาตรฐาน · ภาพซ้อนทับ",
               "และบันทึกเข้าวงจรแพทย์ตรวจทาน"], width=5.6)
    caption(doc, "แผนภาพที่ 2 กล่องเส้นหนาคือด่านบังคับที่ตัดการทำงานของโมดูลวินิจฉัยทั้งหมดเมื่อไม่ผ่าน")
    page_break(doc)


# ---------------------------------------------------------------------------
# Part A divider
# ---------------------------------------------------------------------------

def part_divider(doc, title, subtitle, items):
    p = _para(doc, space_before=120, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, title, size=SZ_TITLE, bold=True)
    p = _para(doc, space_before=0, space_after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, subtitle, size=SZ_SUBTITLE)
    rule(doc, space_before=0, space_after=12)
    for line in items:
        p = _para(doc, space_after=5, indent=1.55)
        _run(p, line, size=SZ_BODY)
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 4 -- fibrosis
# ---------------------------------------------------------------------------

def section_4(doc):
    h1(doc, "4. โมเดลจัดระดับพังผืดตับ F0–F4 (FibrosisNet Ensemble)")
    body(doc, "โมดูลนี้เป็นโครงข่ายประสาทเทียมที่ออกแบบและฝึกขึ้นใหม่ทั้งหมด ไม่ได้ดัดแปลงจากส่วนใด "
              "ของผลงานเดิม จุดที่แตกต่างจากแนวทางทั่วไปคือ โมเดลไม่ได้ถูกออกแบบให้จำแนกระดับพังผืด "
              "เป็นห้าคลาสโดยตรง แต่ถูกออกแบบให้ถดถอยค่าความแข็งของเนื้อตับในสเกลลอการิทึม "
              "แล้วจึงแปลงเป็นระดับในภายหลัง")

    h2(doc, "4.1 หลักคิดเบื้องหลังการออกแบบ")
    body(doc, "ป้ายกำกับระดับพังผืดในชุดข้อมูลไม่ใช่ป้ายที่เป็นอิสระต่อกัน แต่เกิดจากการแบ่งช่วงค่าความแข็งตับ "
              "ที่วัดด้วยเครื่อง TE ตามเกณฑ์คงที่สี่จุด คือ 6.0, 7.1, 8.7 และ 10.3 กิโลพาสคาล "
              "การฝึกโมเดลให้จำแนกคลาสโดยตรงจึงเป็นการทิ้งข้อมูลเชิงปริมาณที่มีอยู่แล้ว และทำให้โมเดล "
              "ไม่ทราบว่าคลาสที่อยู่ติดกันนั้นใกล้กันมากน้อยเพียงใด")
    body(doc, "โมเดลนี้จึงถดถอยค่า log(kPa) เป็นเป้าหมายหลัก โดยมีหัวทำนายเสริมอีกสองหัวช่วยกำกับให้ "
              "การเรียนรู้สอดคล้องกับโครงสร้างลำดับขั้นของโรค สถาปัตยกรรมสามหัวทำนายบนแกนร่วมเดียว "
              "พร้อมเวกเตอร์ฝังมุมตรวจ คือลักษณะทางเทคนิคเฉพาะของโมเดลนี้")

    h2(doc, "4.2 โครงสร้างเครือข่าย")
    table(
        doc,
        ["ส่วนประกอบ", "รายละเอียด"],
        [
            ("แกนสกัดลักษณะเด่น", "ResNet-18 ที่ถอดชั้นจำแนกออก เหลือเวกเตอร์ลักษณะเด่น 512 มิติ\n"
                                  "เริ่มต้นด้วยน้ำหนักที่ฝึกมาก่อนจากชุดภาพทั่วไป"),
            ("เวกเตอร์ฝังมุมตรวจ", "ตารางฝังขนาด 4 ตำแหน่ง × 8 มิติ ครอบคลุมมุมตรวจสามมุม\n"
                                   "บวกตำแหน่งสำรองสำหรับกรณีไม่ทราบมุม"),
            ("การรวมสัญญาณ", "ต่อเวกเตอร์ลักษณะเด่นกับเวกเตอร์ฝังมุมตรวจ ได้ 520 มิติ\n"
                             "แล้วผ่านชั้น dropout อัตรา 0.3"),
            ("หัวทำนายที่ 1", "ชั้นเชิงเส้น 520 → 1 ทำนายค่า log(kPa) เป็นเป้าหมายหลัก"),
            ("หัวทำนายที่ 2", "ชั้นเชิงเส้น 520 → 4 ทำงานแบบ CORN คือแยกเป็นสี่งานย่อย\n"
                              "แต่ละงานตอบว่าเลยระดับที่ k ไปหรือไม่ เมื่อทราบว่าถึงระดับ k แล้ว"),
            ("หัวทำนายที่ 3", "ชั้นเชิงเส้น 520 → 4 ทำนายระดับตามระบบ SWE\n"
                              "เป็นงานเสริมที่ข้ามตัวอย่างซึ่งไม่มีค่าวัด SWE"),
            ("การจัดการช่องสัญญาณ", "ภาพระดับสีเทาถูกทำซ้ำเป็นสามช่องสัญญาณ แทนการเฉลี่ยตัวกรอง\n"
                                    "ชั้นแรก เพื่อรักษาตัวกรองที่ฝึกมาก่อนไว้ครบ"),
        ],
        widths=[1.7, 4.8],
    )

    h2(doc, "4.3 ฟังก์ชันสูญเสียแบบผสมสามส่วน")
    body(doc, "การฝึกใช้ผลรวมถ่วงน้ำหนักของสามองค์ประกอบ โดยให้น้ำหนักหลักกับการถดถอย และให้หัวทำนาย "
              "อีกสองหัวทำหน้าที่กำกับทิศทางการเรียนรู้")
    table(
        doc,
        ["องค์ประกอบ", "รูปแบบ", "น้ำหนัก"],
        [
            ("การถดถอย log(kPa)", "Huber loss ค่าพารามิเตอร์ delta เท่ากับ 0.3", "1.0"),
            ("ลำดับขั้นแบบ CORN", "ผลรวม binary cross-entropy ของสี่งานย่อย\n"
                                  "โดยแต่ละงานคิดเฉพาะตัวอย่างที่ไปถึงระดับนั้น", "0.3"),
            ("ระดับตามระบบ SWE", "cross-entropy ที่ข้ามตัวอย่างซึ่งไม่มีค่าวัด", "0.1"),
        ],
        widths=[1.75, 3.85, 0.9],
        align_cols=[None, None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    lead(doc, "คุณสมบัติที่ได้จากหัวทำนาย CORN: ",
         "การคูณสะสมของค่าซิกมอยด์ตามลำดับงานย่อย ทำให้ความน่าจะเป็นสะสมของการเป็นอย่างน้อยระดับ F1 "
         "ถึง F4 ลดลงตามลำดับเสมอโดยโครงสร้าง จึงไม่เกิดกรณีที่โมเดลบอกว่าโอกาสเป็นอย่างน้อย F3 "
         "สูงกว่าโอกาสเป็นอย่างน้อย F2 ซึ่งเป็นข้อผิดพลาดที่พบได้บ่อยในโมเดลจำแนกหลายคลาสทั่วไป")

    h2(doc, "4.4 การเตรียมภาพและการกำหนด ROI")
    body(doc, "โหมดการเตรียมภาพที่ใช้จริงชื่อ roi_masked_bbox มีขั้นตอนดังนี้")
    numbered(doc, 1, "ทำความสะอาดหน้ากากตับด้วยการปิดช่องว่างด้วยตัวประกอบทรงรีขนาด 15×15 "
                     "เก็บเฉพาะองค์ประกอบเชื่อมต่อที่ใหญ่ที่สุด แล้วเติมรูภายใน "
                     "ขั้นนี้จำเป็นเพราะจากการวัดพบว่าหน้ากากที่ได้จากโครงข่ายมีองค์ประกอบแยกกัน "
                     "โดยมัธยฐาน 2 ชิ้น และมากที่สุดถึง 14 ชิ้น")
    numbered(doc, 2, "คูณภาพระดับสีเทาด้วยหน้ากากที่ทำความสะอาดแล้ว ทำให้พิกเซลนอกเนื้อตับเป็นศูนย์")
    numbered(doc, 3, "ครอบตัดตามกรอบสี่เหลี่ยมของหน้ากาก โดยขยายกรอบออกด้านละ 5 เปอร์เซ็นต์")
    numbered(doc, 4, "ปรับขนาดเป็นสี่เหลี่ยมจัตุรัส 256×256 โดยไม่รักษาอัตราส่วนภาพเดิม")
    lead(doc, "เหตุผลของการทำลายอัตราส่วนภาพโดยตั้งใจ: ",
         "การตรวจสอบชุดข้อมูลพบว่าความละเอียดของภาพพัวพันกับป้ายกำกับอย่างมีนัยสำคัญ "
         "ภาพขนาด 720×1000 มีค่าความแข็งเฉลี่ย 5.10 กิโลพาสคาล และไม่มีตัวอย่างระดับ F4 เลย "
         "ขณะที่ภาพขนาด 730×1020 มีค่าเฉลี่ย 6.75 กิโลพาสคาล และมีสัดส่วน F4 ถึง 11.9 เปอร์เซ็นต์ "
         "หากคงอัตราส่วนภาพไว้ โมเดลจะสามารถทำนายระดับพังผืดได้จากรูปทรงของภาพแทนที่จะดูเนื้อตับ "
         "การยืดภาพเป็นจัตุรัสจึงเป็นมาตรการตัดทางลัดของการเรียนรู้ที่ตั้งใจออกแบบไว้")
    body(doc, "นอกจากนี้ ระบบยังนิยามโหมดเตรียมภาพสำหรับใช้เป็นชุดควบคุมเชิงลบไว้ด้วย ได้แก่ โหมดที่ใช้ "
              "เฉพาะพิกเซลนอกกรวยสัญญาณ และโหมดที่ใช้เฉพาะรูปทรงของหน้ากากโดยไม่มีเนื้อภาพ "
              "ทั้งสองโหมดนี้ใช้พิสูจน์ว่าโมเดลเรียนรู้จากลักษณะเนื้อตับจริง ดังผลในหัวข้อ 4.9")
    page_break(doc)

    h2(doc, "4.5 การเพิ่มความหลากหลายของข้อมูลที่ปรับเฉพาะภาพอัลตราซาวด์")
    table(
        doc,
        ["วิธี", "ค่าที่ใช้", "เหตุผล"],
        [
            ("ครอบตัดและปรับขนาดแบบสุ่ม", "สัดส่วนพื้นที่ 0.6–1.0\nอัตราส่วน 0.8–1.25",
             "ทำลายร่องรอยของความละเอียดและการจัดวาง\nที่ยังเหลืออยู่"),
            ("พลิกภาพซ้ายขวา", "ความน่าจะเป็น 0.5", "ทิศทางซ้ายขวาไม่มีความหมายทางกายวิภาค"),
            ("ไม่พลิกภาพบนล่าง", "ปิดการใช้งาน",
             "แกนตั้งของภาพคือทิศทางความลึกของลำคลื่น\nการพลิกทำให้ได้ภาพที่เป็นไปไม่ได้ทางกายภาพ"),
            ("แปลงแอฟฟีน", "หมุน 7 องศา · เลื่อน 0.05\n· เฉือน 5 องศา", "จำลองการวางหัวตรวจที่ต่างกันเล็กน้อย"),
            ("ปรับความสว่างและคอนทราสต์", "อย่างละ 0.25", "จำลองการตั้งค่าเครื่องที่ต่างกัน"),
            ("เบลอแบบเกาส์เซียน", "ความน่าจะเป็น 0.2", "จำลองภาพที่โฟกัสไม่คม"),
            ("สัญญาณรบกวน speckle", "แบบคูณ ค่าเบี่ยงเบนสูงสุด 0.03\nความน่าจะเป็น 0.3",
             "จุดรบกวนในภาพอัลตราซาวด์เป็นสัญญาณรบกวน\nแบบคูณ ไม่ใช่แบบบวก จึงต้องจำลองให้ตรงลักษณะ"),
            ("ลบบางส่วนของภาพแบบสุ่ม", "ความน่าจะเป็น 0.25\nพื้นที่ 0.02–0.1",
             "บังคับให้โมเดลไม่พึ่งพาบริเวณเดียว"),
        ],
        widths=[1.75, 1.85, 2.9],
    )
    body(doc, "ขั้นตอนตอนอนุมานใช้เพียงการปรับขนาดเป็น 256×256 และการปรับมาตรฐานค่าพิกเซลเท่านั้น "
              "ไม่มีการเพิ่มความหลากหลายใด ๆ")

    h2(doc, "4.6 การจัดกลุ่มครั้งตรวจและการแบ่งข้อมูลแบบไม่รั่วไหล")
    body(doc, "ชุดข้อมูลที่ใช้ฝึกมีรหัสประจำตัวที่เป็นรหัสรอบการเก็บข้อมูล ไม่ใช่รหัสผู้ป่วยที่แท้จริง "
              "การแบ่งข้อมูลตามรหัสนั้นตรง ๆ จึงยังเสี่ยงต่อการรั่วไหล ระบบจึงสร้างหน่วยวิเคราะห์ใหม่ "
              "เรียกว่าครั้งตรวจ ด้วยวิธีต่อไปนี้")
    numbered(doc, 1, "อ่านเวลาที่บันทึกภาพจากรหัสไฟล์ โดยแปลงเลขฐานสิบหกสี่ไบต์แรกของรหัสเป็นเวลา")
    numbered(doc, 2, "จัดภาพที่อยู่ในรหัสเดียวกันและมีระยะห่างเวลาติดกันไม่เกิน 300 วินาที "
                     "ให้เป็นครั้งตรวจเดียวกัน และเริ่มครั้งตรวจใหม่เมื่อเปลี่ยนรหัสหรือระยะห่างเกินเกณฑ์")
    numbered(doc, 3, "แบ่งย่อยกลุ่มที่ได้ด้วยค่าความแข็งตับอีกชั้นหนึ่ง เนื่องจากผู้ป่วยสองรายอาจถูกตรวจ "
                     "ต่อเนื่องกันจนไม่มีช่องว่างเวลาให้แยกได้ แต่ย่อมมีค่าความแข็งตับต่างกัน")
    body(doc, "ผลคือแปลงภาพ 1,772 ภาพให้เป็นครั้งตรวจ 730 ครั้ง กระจายเป็นระดับ F0 จำนวน 497 ครั้ง "
              "F1 จำนวน 105 ครั้ง F2 จำนวน 48 ครั้ง F3 จำนวน 37 ครั้ง และ F4 จำนวน 43 ครั้ง")
    body(doc, "การแบ่งข้อมูลใช้วิธี StratifiedGroupKFold ที่จัดกลุ่มตามรหัสประจำตัวและถ่วงสัดส่วนระดับ "
              "ของครั้งตรวจไปพร้อมกัน กำหนดเป็น 5 ส่วนต่อรอบ ทำซ้ำ 3 รอบ รวมทั้งสิ้น 15 ชุดการแบ่ง "
              "และตรึงผลการแบ่งไว้เป็นไฟล์เพื่อไม่ให้เปลี่ยนระหว่างการทดลอง "
              "ภายในชุดฝึกของแต่ละชุดยังแบ่งย่อยอีก 4 ส่วน สำหรับใช้หยุดการฝึกและปรับเทียบเกณฑ์")
    lead(doc, "การตรวจสอบที่บังคับใช้: ",
         "ก่อนเริ่มฝึก ระบบตรวจว่าไม่มีครั้งตรวจใดปรากฏซ้ำข้ามชุด ทุกครั้งตรวจถูกใช้ครบ "
         "ไม่มีรหัสประจำตัวใดรั่วข้ามชุด และทุกชุดตรวจสอบต้องมีครั้งตรวจระดับ F3 อย่างน้อย 5 ครั้ง "
         "และระดับ F4 อย่างน้อย 6 ครั้ง หากเงื่อนไขใดไม่ผ่าน การฝึกจะไม่เริ่ม")

    h2(doc, "4.7 ขั้นตอนการฝึกและการรวมเป็นคณะโมเดล")
    table(
        doc,
        ["พารามิเตอร์", "ค่า", "พารามิเตอร์", "ค่า"],
        [
            ("จำนวนรอบสูงสุด", "40", "อัตราเรียนรู้ของหัวทำนาย", "3×10⁻⁴"),
            ("รอบอดทนก่อนหยุด", "8", "อัตราเรียนรู้ของแกนสกัด", "3×10⁻⁵"),
            ("ขนาดชุดย่อย", "32", "ตัวปรับพารามิเตอร์", "AdamW"),
            ("ขนาดภาพเข้า", "256×256", "ค่าลดทอนน้ำหนัก", "0.01"),
            ("รอบอุ่นเครื่อง", "3", "การตัดขนาดเกรเดียนต์", "1.0"),
            ("อัตราเฉลี่ยเคลื่อนที่", "0.999", "อัตรา dropout", "0.3"),
        ],
        widths=[1.75, 1.5, 1.85, 1.4],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER, None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    bullet(doc, "อัตราเรียนรู้เพิ่มขึ้นเชิงเส้นในช่วงอุ่นเครื่อง จากนั้นลดลงตามฟังก์ชันโคไซน์",
           label="ตารางอัตราเรียนรู้: ")
    bullet(doc, "ใช้ค่าเฉลี่ยเคลื่อนที่ของน้ำหนักที่มีการอุ่นค่าอัตราลดทอนตามจำนวนก้าว "
                "เนื่องจากหนึ่งรอบการฝึกมีเพียงประมาณ 35 ก้าว หากใช้อัตราคงที่ 0.999 ตั้งแต่ต้น "
                "ค่าเฉลี่ยจะยังคงมีน้ำหนักเริ่มต้นแบบสุ่มปนอยู่มากแม้ฝึกครบ 40 รอบแล้ว",
           label="การเฉลี่ยน้ำหนัก: ")
    bullet(doc, "ใช้ค่า AUROC ระดับครั้งตรวจสำหรับการแยกอย่างน้อยระดับ F2 บนชุดแบ่งย่อยภายใน "
                "และคำนวณจากน้ำหนักที่ผ่านการเฉลี่ยแล้ว", label="เกณฑ์หยุดการฝึก: ")
    bullet(doc, "ทำนายสองครั้งคือภาพปกติและภาพพลิกซ้ายขวา แล้วเฉลี่ยอย่างละครึ่ง "
                "ทั้งในสเกล log(kPa) และในสเกลความน่าจะเป็นสะสม", label="การทำนายซ้ำตอนอนุมาน: ")
    body(doc, "คณะโมเดลที่นำไปใช้งานจริงประกอบด้วยโมเดลห้าตัว ซึ่งเป็นห้าชุดการแบ่งของรอบทำซ้ำที่หนึ่ง "
              "ทุกตัวใช้แกนสกัด ResNet-18 และโหมดเตรียมภาพ roi_masked_bbox เหมือนกัน")
    page_break(doc)

    h2(doc, "4.8 การอนุมานและการแปลงเป็นระดับ")
    body(doc, "ขั้นตอนการอนุมานมีลำดับดังนี้")
    numbered(doc, 1, "เตรียม ROI ตามหัวข้อ 4.4 แปลงเป็นเทนเซอร์ ทำซ้ำเป็นสามช่องสัญญาณ "
                     "และปรับมาตรฐานด้วยค่าเฉลี่ยและส่วนเบี่ยงเบนมาตรฐานชุดเดียวกับที่ใช้ตอนฝึก")
    numbered(doc, 2, "ให้โมเดลทั้งห้าตัวทำนาย โดยแต่ละตัวเฉลี่ยผลของภาพปกติกับภาพพลิกซ้ายขวา")
    numbered(doc, 3, "เฉลี่ยค่า log(kPa) ของทั้งห้าตัว แล้วแปลงกลับด้วยฟังก์ชันเลขชี้กำลัง "
                     "ได้ค่าประมาณความแข็งตับเป็นกิโลพาสคาล โดยไม่มีการคูณสเกล บวกค่าชดเชย "
                     "หรือเปิดตารางเทียบใด ๆ เพิ่มเติม")
    numbered(doc, 4, "เฉลี่ยความน่าจะเป็นสะสมของทั้งห้าตัว ได้ค่าโอกาสเป็นอย่างน้อยระดับ F2 "
                     "อย่างน้อยระดับ F3 และเป็นระดับ F4")
    numbered(doc, 5, "แปลงเป็นระดับสองทางขนานกัน ทางแรกใช้เกณฑ์ทางคลินิกคงที่ ทางที่สองใช้เกณฑ์ "
                     "ที่ปรับเทียบตามความชุกซึ่งอธิบายในหัวข้อ 4.9")
    numbered(doc, 6, "คำนวณระดับความเสี่ยงจากค่าโอกาสเป็นอย่างน้อยระดับ F2 และคืนกรอบสี่เหลี่ยม "
                     "ของ ROI เป็นพิกัดในภาพต้นฉบับ พร้อมค่าการกระจายของคณะโมเดล "
                     "ซึ่งคือผลต่างระหว่างค่าสูงสุดกับค่าต่ำสุดที่โมเดลทั้งห้าทำนาย")

    h2(doc, "4.9 การปรับเทียบเกณฑ์และระดับความเสี่ยง")
    body(doc, "ระบบไม่ได้ใช้การปรับเทียบด้วยอุณหภูมิ การปรับเทียบแบบ Platt หรือการถดถอยแบบเอกโทน "
              "แต่ใช้วิธีจับคู่ควอนไทล์กับความชุกที่สังเกตได้ กล่าวคือ สะสมสัดส่วนของแต่ละระดับในชุดปรับเทียบ "
              "แล้วกำหนดจุดตัดที่ควอนไทล์เดียวกันของคะแนนที่โมเดลทำนาย")
    lead(doc, "เหตุผล: ",
         "แบบจำลองถดถอยย่อมหดค่าทำนายเข้าหาค่ากลางเสมอ หากนำค่าที่หดแล้วไปเทียบกับเกณฑ์ทางคลินิก "
         "ซึ่งกำหนดบนสเกลจริง ระบบจะประเมินระดับสูง ๆ ต่ำกว่าความเป็นจริงอย่างเป็นระบบ "
         "การปรับเทียบตามความชุกจึงย้ายจุดตัดมาไว้บนสเกลของคะแนนที่โมเดลผลิตจริง")
    table(
        doc,
        ["จุดตัด", "ค่าในสเกล log", "ค่าเทียบเท่าเป็น kPa", "เกณฑ์ทางคลินิก TE"],
        [
            ("F0 / F1", "1.3303", "3.78", "6.0"),
            ("F1 / F2", "1.4456", "4.24", "7.1"),
            ("F2 / F3", "1.5376", "4.65", "8.7"),
            ("F3 / F4", "1.6352", "5.13", "10.3"),
        ],
        widths=[1.3, 1.6, 1.9, 1.7],
        align_cols=[WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    body(doc, "ความห่างระหว่างสองคอลัมน์ขวาแสดงระดับการหดตัวของค่าทำนายอย่างชัดเจน และเป็นเหตุผล "
              "ที่ระบบต้องรายงานทั้งค่าประมาณและข้อความกำกับข้อจำกัดควบคู่กันเสมอ")
    body(doc, "ระดับความเสี่ยงคำนวณจากค่าโอกาสเป็นอย่างน้อยระดับ F2 โดยใช้จุดแบ่งที่ 0.15 และ 0.30 "
              "และแต่ละระดับมาพร้อมอัตราที่สังเกตได้จริงจากครั้งตรวจนอกชุดฝึก ดังตาราง")
    table(
        doc,
        ["ระดับความเสี่ยง", "ช่วงค่าโอกาส ≥F2", "จำนวนครั้งตรวจ", "อัตราที่พบ ≥F2", "อัตราที่พบ F4"],
        [
            ("ต่ำ", "น้อยกว่า 0.15", "475", "9.68 %", "2.32 %"),
            ("ปานกลาง", "0.15 ถึง 0.30", "181", "24.31 %", "5.52 %"),
            ("สูง", "ตั้งแต่ 0.30", "74", "51.35 %", "29.73 %"),
        ],
        widths=[1.35, 1.5, 1.3, 1.2, 1.15],
        align_cols=[None] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    body(doc, "การรายงานอัตราที่สังเกตได้จริงควบคู่กับระดับความเสี่ยง ทำให้ผู้ใช้ทราบว่าคำว่าเสี่ยงสูง "
              "ในระบบนี้หมายถึงโอกาสพบพังผืดตั้งแต่ระดับ F2 ประมาณครึ่งหนึ่ง ไม่ใช่ความแน่นอน")
    page_break(doc)

    h2(doc, "4.10 ชั้นตัดสินระดับที่รายงานต่อผู้ใช้")
    body(doc, "ก่อนแสดงผล ระบบมีชั้นตัดสินอีกชั้นหนึ่งที่ไม่ใช้ระดับที่โมเดลให้มาโดยตรง แต่คำนวณระดับใหม่ "
              "จากค่าความน่าจะเป็นและค่าประมาณความแข็งตับ ด้วยเกณฑ์สัมบูรณ์ที่เรียงลำดับจากรุนแรงไปเบา "
              "และหยุดที่เงื่อนไขแรกที่เป็นจริง")
    table(
        doc,
        ["ลำดับ", "เงื่อนไข", "ระดับที่ให้", "ระดับความเสี่ยง"],
        [
            ("1", "โอกาสเป็น F4 ตั้งแต่ 0.25\nหรือ (kPa ตั้งแต่ 6.0 และโอกาส ≥F3 ตั้งแต่ 0.25)", "F4", "สูง"),
            ("2", "โอกาส ≥F3 ตั้งแต่ 0.40 หรือ kPa ตั้งแต่ 5.5", "F3", "สูง"),
            ("3", "โอกาส ≥F2 ตั้งแต่ 0.35 หรือ kPa ตั้งแต่ 4.6", "F2", "ปานกลาง"),
            ("4", "โอกาส ≥F2 ตั้งแต่ 0.25 หรือ kPa ตั้งแต่ 4.0", "F1", "ต่ำ"),
            ("5", "ไม่เข้าเงื่อนไขข้างต้น", "F0", "ต่ำ"),
        ],
        widths=[0.6, 3.6, 1.15, 1.15],
        align_cols=[WD_ALIGN_PARAGRAPH.CENTER, None,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    body(doc, "ค่าความมั่นใจที่รายงานถูกจำกัดช่วงตามระดับที่ได้ คือ ระดับ F3 และ F4 อยู่ในช่วง 0.78 ถึง 0.95 "
              "ระดับ F2 อยู่ในช่วง 0.72 ถึง 0.90 ระดับ F1 อยู่ในช่วง 0.75 ถึง 0.90 "
              "และระดับ F0 อยู่ในช่วง 0.85 ถึง 0.96 โดยคำนวณจากส่วนเติมเต็มของโอกาสเป็นอย่างน้อยระดับ F2")

    h2(doc, "4.11 การแก้ปัญหาการประเมินเกินจริงในภาพตับปกติ")
    body(doc, "ชั้นตัดสินรุ่นก่อนหน้าทำงานแบบไล่ระดับขึ้นทางเดียว คือเริ่มจากระดับที่โมเดลให้มา "
              "แล้วเลื่อนขึ้นเมื่อเข้าเงื่อนไขบางข้อ แต่ไม่มีทางเลื่อนลงได้เลย ผลคือภาพตับปกติที่ได้ระดับ "
              "ความเสี่ยงปานกลางหรือสูง จะถูกดันขึ้นเป็นระดับ F2 ถึง F4 ทันทีโดยไม่ต้องตรวจสอบว่า "
              "ค่าโอกาสของระดับนั้นสูงจริงหรือไม่")
    table(
        doc,
        ["ประเด็น", "รุ่นก่อนแก้ไข", "รุ่นปัจจุบัน"],
        [
            ("กลไกหลัก", "ไล่ระดับขึ้นจากระดับที่โมเดลให้มา\nเลื่อนลงไม่ได้",
             "คำนวณใหม่จากเกณฑ์สัมบูรณ์\nทุกครั้ง"),
            ("เส้นทางสู่ระดับ F4", "โอกาสเป็น F4 ตั้งแต่ 0.30\nหรือถูกดันขึ้นจากระดับความเสี่ยงสูง",
             "โอกาสเป็น F4 ตั้งแต่ 0.25\nหรือเงื่อนไขร่วมของ kPa กับโอกาส ≥F3"),
            ("เส้นทางสู่ระดับ F3", "โอกาส ≥F3 ตั้งแต่ 0.45\nหรือถูกดันขึ้นจากระดับความเสี่ยง",
             "โอกาส ≥F3 ตั้งแต่ 0.40 หรือ kPa ตั้งแต่ 5.5\nและเป็นเส้นทางเดียวเท่านั้น"),
            ("แถบระดับ F1", "ไม่มี ภาพที่คาบเกี่ยวถูกดันขึ้น F2", "มี เป็นพื้นที่รองรับภาพที่คาบเกี่ยว"),
            ("ค่าความมั่นใจของ F0", "ช่วง 0.78 ถึง 0.95", "ช่วง 0.85 ถึง 0.96"),
        ],
        widths=[1.35, 2.6, 2.55],
    )
    lead(doc, "ผลเชิงพฤติกรรม: ",
         "การเปลี่ยนแปลงนี้ทำให้เงื่อนไขการยกระดับทุกเส้นทางต้องอ้างอิงค่าความน่าจะเป็นของระดับนั้นโดยตรง "
         "ภาพตับปกติที่ได้ค่าโอกาสไม่ถึงเกณฑ์จึงคงอยู่ที่ระดับ F0 หรือ F1 และการยกพื้นค่าความมั่นใจของ "
         "ระดับ F0 ขึ้นสะท้อนว่าระบบยืนยันผลปกติหนักแน่นขึ้นด้วย")
    page_break(doc)

    h2(doc, "4.12 ผลการวัดประสิทธิภาพและการควบคุมเชิงลบ")
    body(doc, "การวัดผลใช้การตรวจสอบไขว้แบบจัดกลุ่มครบ 15 ชุด บนครั้งตรวจ 730 ครั้ง "
              "รายงานเป็นค่า AUROC ระดับครั้งตรวจ")
    table(
        doc,
        ["ชุดการทดลอง", "ลักษณะ", "AUROC ≥F2", "≥F3", "F4"],
        [
            ("โมเดลที่ใช้งานจริง", "ResNet-18 บน ROI เนื้อตับ", "0.7180 ± 0.0477", "0.7455", "0.7645"),
            ("ฐานเปรียบเทียบเชิงพื้นผิว", "ลักษณะเด่นเชิงสถิติ + GBM", "0.6859 ± 0.0700", "0.7049", "0.7639"),
            ("ควบคุมเชิงลบที่ 1", "ใช้เฉพาะรูปทรงหน้ากาก\nไม่มีเนื้อภาพ", "0.6261", "0.6377", "0.6764"),
            ("ควบคุมเชิงลบที่ 2", "ใช้เฉพาะพิกเซลนอกกรวยสัญญาณ", "0.6072", "0.6051", "0.6154"),
            ("ควบคุมเชิงลบที่ 3", "ใช้เฉพาะข้อมูลกำกับ ไม่มีพิกเซล", "0.5934", "0.5734", "0.6065"),
            ("ฐานทายคลาสส่วนมาก", "ไม่ใช้ข้อมูลใดเลย", "0.5000", "0.5000", "0.5000"),
        ],
        widths=[1.55, 2.1, 1.3, 0.75, 0.8],
        align_cols=[None, None] + [WD_ALIGN_PARAGRAPH.CENTER] * 3,
    )
    body(doc, "การทดสอบนัยสำคัญใช้การสุ่มซ้ำแบบ bootstrap ระดับรหัสประจำตัว จำนวน 2,000 รอบ "
              "บนคะแนนที่ปรับมาตรฐานรายชุดการแบ่ง ค่า AUROC รวมของโมเดลที่ใช้งานจริงเท่ากับ 0.7386 "
              "และผลต่างจากชุดควบคุมเชิงลบทั้งสามชุดมีช่วงความเชื่อมั่นที่ไม่คร่อมศูนย์ทั้งหมด ดังนี้")
    bullet(doc, "สูงกว่าชุดที่ใช้พิกเซลนอกกรวยสัญญาณ 0.0967 ช่วงความเชื่อมั่น 0.0404 ถึง 0.1559")
    bullet(doc, "สูงกว่าชุดที่ใช้เฉพาะรูปทรงหน้ากาก 0.1010 ช่วงความเชื่อมั่น 0.0504 ถึง 0.1569")
    bullet(doc, "สูงกว่าชุดที่ใช้เฉพาะข้อมูลกำกับ 0.1474 ช่วงความเชื่อมั่น 0.0781 ถึง 0.2151")
    body(doc, "ผลนี้ยืนยันว่าโมเดลเรียนรู้จากลักษณะของเนื้อตับจริง ไม่ได้อาศัยรูปทรงของหน้ากาก "
              "ตัวอักษรที่เครื่องพิมพ์ทับ หรือข้อมูลกำกับของไฟล์ ซึ่งเป็นประเด็นที่ต้องพิสูจน์ "
              "เมื่อกล่าวอ้างว่าโมเดลตรวจจับพยาธิสภาพจากภาพ")

    h2(doc, "4.13 ข้อจำกัดที่ระบบแจ้งผู้ใช้เสมอ")
    body(doc, "ระบบแนบข้อความกำกับข้อจำกัดไปกับผลของโมดูลนี้ทุกครั้ง เนื้อหาอ้างอิงตัวเลขที่วัดได้จริง ดังนี้")
    bullet(doc, "ค่าความแข็งตับที่ทำนายได้มีช่วงการกระจายเพียงประมาณ 39 เปอร์เซ็นต์ของช่วงจริง")
    bullet(doc, "ครั้งตรวจที่เป็นระดับ F4 จริงมีค่าเฉลี่ยที่วัดได้ 15.09 กิโลพาสคาล "
                "แต่โมเดลทำนายเฉลี่ยประมาณ 5.97 กิโลพาสคาล")
    bullet(doc, "ค่าสัมประสิทธิ์ความสอดคล้องแบบถ่วงกำลังสองเท่ากับ 0.37 บนครั้งตรวจ 730 ครั้ง")
    bullet(doc, "ความไวต่อการตรวจพบตับแข็งอยู่ที่ประมาณ 35 เปอร์เซ็นต์")
    body(doc, "การเปิดเผยข้อจำกัดเหล่านี้เป็นส่วนหนึ่งของการออกแบบ ระบบจึงวางตำแหน่งค่าประมาณ kPa "
              "เป็นค่าประกอบการพิจารณา ไม่ใช่ค่าทดแทนการวัดด้วยเครื่องมือมาตรฐาน "
              "และให้น้ำหนักกับการแบ่งระดับความเสี่ยงมากกว่าค่าตัวเลขเดี่ยว")
    page_break(doc)

    h2(doc, "4.14 แผนภาพที่ 3 โครงสร้าง FibrosisNet")
    dbox(doc, ["ข้อมูลเข้า", "ภาพระดับสีเทา + หน้ากากตับ + รหัสมุมตรวจ"], width=5.4)
    darrow(doc)
    dbox(doc, ["การเตรียม ROI", "ทำความสะอาดหน้ากาก → คูณภาพด้วยหน้ากาก → ครอบตัดกรอบ + 5%",
               "→ ยืดเป็น 256×256 โดยไม่รักษาอัตราส่วน"], width=5.4)
    darrow(doc)
    drow(
        doc,
        [
            ["แกนสกัด ResNet-18", "เวกเตอร์ 512 มิติ"],
            ["ตารางฝังมุมตรวจ", "4 ตำแหน่ง × 8 มิติ"],
        ],
        widths=[3.0, 2.4],
    )
    darrow(doc, "ต่อกันเป็น 520 มิติ แล้วผ่าน dropout 0.3")
    drow(
        doc,
        [
            ["หัวที่ 1  ถดถอย", "520 → 1", "ค่า log(kPa)", "น้ำหนัก 1.0"],
            ["หัวที่ 2  CORN", "520 → 4", "โอกาสสะสม ≥F1..≥F4", "น้ำหนัก 0.3"],
            ["หัวที่ 3  SWE", "520 → 4", "ระดับตามระบบ SWE", "น้ำหนัก 0.1"],
        ],
        widths=[2.2, 2.2, 2.1],
    )
    darrow(doc, "เฉลี่ยผลของคณะโมเดล 5 ตัว พร้อมการทำนายซ้ำแบบพลิกซ้ายขวา")
    dbox(doc, ["ผลลัพธ์", "kPa = exp(ค่าเฉลี่ย log(kPa)) · โอกาส ≥F2, ≥F3, F4",
               "ระดับความเสี่ยงจากจุดแบ่ง 0.15 และ 0.30 · กรอบ ROI"], width=5.4)
    darrow(doc)
    dbox(doc, ["ชั้นตัดสินระดับที่รายงาน", "เกณฑ์สัมบูรณ์ 5 ลำดับ ตามตารางในหัวข้อ 4.10"],
         width=5.4, heavy=True)
    caption(doc, "แผนภาพที่ 3 หัวทำนายทั้งสามใช้แกนสกัดร่วมกัน จึงเรียนรู้ลักษณะเด่นชุดเดียว "
                 "ที่ต้องตอบได้ทั้งค่าต่อเนื่องและลำดับขั้นไปพร้อมกัน")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 5 -- steatosis
# ---------------------------------------------------------------------------

def section_5(doc):
    h1(doc, "5. โมดูลประเมินไขมันพอกตับ S0–S3")

    h2(doc, "5.1 ลักษณะทางเทคนิคของโมดูล")
    lead(doc, "ข้อเท็จจริงที่ต้องระบุให้ชัด: ",
         "โมดูลนี้ไม่ใช่โครงข่ายประสาทเทียม และไม่มีไฟล์น้ำหนักโมเดลใด ๆ ในระบบ "
         "แต่เป็นการคำนวณเชิงฟิสิกส์แบบกำหนดได้ ที่วัดค่าทางแสงเสียงจากพิกเซลของเนื้อตับโดยตรง "
         "ให้ผลลัพธ์เหมือนเดิมทุกครั้งเมื่อป้อนภาพเดิม")
    body(doc, "การจัดหมวดหมู่นี้สำคัญต่อการกำหนดแนวทางขอความคุ้มครอง เพราะสาระของสิ่งที่สร้างขึ้นใหม่ "
              "อยู่ที่วิธีกำหนดบริเวณสุ่มตัวอย่าง วิธีคำนวณค่าที่ทนต่อสิ่งรบกวน และกฎรวมคะแนน "
              "ไม่ได้อยู่ที่พารามิเตอร์ที่ได้จากการเรียนรู้")
    body(doc, "หลักการทางกายภาพเบื้องหลังคือ เมื่อเนื้อตับมีไขมันสะสมมากขึ้น ลำคลื่นเสียงจะถูกลดทอน "
              "ระหว่างเดินทางลงลึกมากกว่าปกติ ทำให้บริเวณตื้นสว่างขึ้นเมื่อเทียบกับบริเวณลึก "
              "และเนื้อตับโดยรวมมีความสว่างสูงขึ้น โมดูลนี้จึงวัดสองสิ่งนี้เป็นตัวชี้วัดหลัก")

    h2(doc, "5.2 การกำหนดบริเวณสุ่มตัวอย่างแบบแบ่งแถบความลึกภายในหน้ากากตับ")
    body(doc, "โมดูลไม่ใช้กรอบคงที่บนเฟรม แต่กำหนดแถบความลึกโดยอ้างอิงกับช่วงความสูงของหน้ากากตับเอง "
              "แล้วตัดทับด้วยหน้ากากอีกชั้นหนึ่ง เพื่อให้ทุกพิกเซลที่นำมาคำนวณเป็นเนื้อตับแน่นอน")
    table(
        doc,
        ["แถบ", "ช่วงความลึกเทียบกับความสูงของตับ", "บทบาท"],
        [
            ("แถบตื้น", "15 % ถึง 45 %",
             "วัดความสว่างของเนื้อตับส่วนต้น\nโดยเว้นยอดกรวยสัญญาณที่มักมีสิ่งแปลกปลอม"),
            ("แถบลึก", "55 % ถึง 90 %",
             "วัดความสว่างของเนื้อตับส่วนลึก\nโดยเว้นแถบล่างสุดที่สัญญาณมักหายไป"),
        ],
        widths=[1.15, 2.4, 2.95],
    )
    lead(doc, "คุณสมบัติสำคัญ: ",
         "เนื่องจากแถบทั้งสองนิยามเป็นสัดส่วนของหน้ากากตับ ไม่ใช่ของเฟรม ค่าที่วัดได้จึงไม่ผันแปร "
         "ตามการตั้งความลึกของเครื่องตรวจ หรือตามขนาดของตับที่ปรากฏในภาพ "
         "หากช่วงความสูงของตับน้อยกว่าหรือเท่ากับ 30 พิกเซล ระบบถือว่าแบ่งแถบไม่ได้ "
         "และใช้ค่าอัตราการลดทอนตั้งต้นที่ 1.05 แทน")

    h2(doc, "5.3 ตัวประมาณค่าเนื้อตับแบบทนต่อสิ่งรบกวน")
    body(doc, "ค่าเฉลี่ยธรรมดามีปัญหาในภาพอัลตราซาวด์ตับ เพราะภายในเนื้อตับมีหลอดเลือดซึ่งเป็นรูมืดเกือบดำ "
              "และผนังหลอดเลือดกับหินปูนซึ่งสว่างจัด ทั้งสองอย่างดึงค่าเฉลี่ยให้เพี้ยน "
              "โมดูลจึงใช้ตัวประมาณที่ตัดปลายทั้งสองข้างก่อนเฉลี่ย")
    numbered(doc, 1, "หากมีพิกเซลน้อยกว่า 10 จุด ให้ใช้ค่าเฉลี่ยธรรมดา หรือค่าตั้งต้น 90 เมื่อไม่มีพิกเซลเลย")
    numbered(doc, 2, "มิฉะนั้น คำนวณควอนไทล์ที่ 10 และที่ 90 ของค่าความสว่าง")
    numbered(doc, 3, "เก็บเฉพาะพิกเซลที่อยู่ระหว่างสองควอนไทล์นั้น แล้วหาค่าเฉลี่ย")
    body(doc, "ชุดทดสอบของโครงการยืนยันคุณสมบัตินี้ไว้โดยตรง คือเมื่อฝังรูมืดขนาด 30×30 พิกเซล "
              "จำนวนสองจุดลงในภาพตับจำลอง ค่าอัตราการลดทอนที่คำนวณได้ต้องเปลี่ยนแปลงน้อยกว่า 0.25")

    h2(doc, "5.4 ตัวชี้วัดสองตัวและสูตรคำนวณ")
    table(
        doc,
        ["ตัวชี้วัด", "สูตร", "ความหมาย"],
        [
            ("อัตราการลดทอนลำคลื่น",
             "A = ค่าเฉลี่ยตัดปลายของแถบตื้น\n÷ ค่าเฉลี่ยตัดปลายของแถบลึก\n(ตัวหารมีค่าต่ำสุดที่ 1.0)",
             "ยิ่งมากยิ่งบ่งชี้ว่าลำคลื่นถูกลดทอน\nระหว่างเดินทางลงลึกมากผิดปกติ"),
            ("ดัชนีความสว่างเนื้อตับ",
             "E = ค่าเฉลี่ยตัดปลาย\nของพิกเซลทั้งหมดในหน้ากากตับ\n(สเกล 0 ถึง 255)",
             "ยิ่งมากยิ่งบ่งชี้ว่าเนื้อตับสว่างกว่าปกติ\nซึ่งสัมพันธ์กับปริมาณไขมัน"),
        ],
        widths=[1.55, 2.5, 2.45],
    )
    body(doc, "นอกจากสองตัวชี้วัดนี้ โมดูลยังรับสัญญาณตัวที่สามจากตัวตรวจจับรอยโรค คือการพบการเปลี่ยนแปลง "
              "ไขมันเฉพาะที่ หรือการเว้นไขมันเฉพาะที่ ซึ่งอธิบายไว้ในหัวข้อ 5.6")

    h2(doc, "5.5 การให้คะแนนแบบบวกและเกณฑ์แบ่งระดับ")
    body(doc, "คะแนนเริ่มต้นที่ศูนย์ แล้วบวกเพิ่มตามเงื่อนไขสามกลุ่ม โดยแต่ละกลุ่มให้คะแนนได้เพียงชั้นเดียว "
              "คือชั้นที่สูงที่สุดที่เข้าเงื่อนไข คะแนนสูงสุดที่เป็นไปได้คือ 5.6")
    table(
        doc,
        ["องค์ประกอบ", "เงื่อนไข", "คะแนนที่บวก"],
        [
            ("อัตราการลดทอน", "ตั้งแต่ 1.45", "+2.2"),
            ("", "ตั้งแต่ 1.30 แต่ไม่ถึง 1.45", "+1.5"),
            ("", "ตั้งแต่ 1.18 แต่ไม่ถึง 1.30", "+0.8"),
            ("ความสว่างเนื้อตับ", "ตั้งแต่ 120", "+1.6"),
            ("", "ตั้งแต่ 95 แต่ไม่ถึง 120", "+0.9"),
            ("", "ตั้งแต่ 80 แต่ไม่ถึง 95", "+0.3"),
            ("ไขมันเฉพาะที่", "พบการเปลี่ยนแปลงหรือการเว้นไขมันเฉพาะที่", "+1.8"),
        ],
        widths=[1.75, 3.55, 1.2],
        align_cols=[None, None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    table(
        doc,
        ["คะแนนรวม", "ระดับ", "ค่าความมั่นใจที่รายงาน"],
        [
            ("ตั้งแต่ 2.6", "S3", "0.88"),
            ("ตั้งแต่ 1.6 แต่ไม่ถึง 2.6", "S2", "0.85"),
            ("ตั้งแต่ 0.8 แต่ไม่ถึง 1.6", "S1", "0.82"),
            ("น้อยกว่า 0.8", "S0", "0.90"),
        ],
        widths=[2.6, 1.6, 2.3],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    body(doc, "ค่าความมั่นใจของโมดูลนี้เป็นค่าคงที่ประจำระดับ ไม่ได้คำนวณจากตัวชี้วัด "
              "เนื่องจากโมดูลไม่มีแบบจำลองความน่าจะเป็นอยู่เบื้องหลัง")
    page_break(doc)

    h2(doc, "5.6 การรับสัญญาณจากตัวตรวจจับรอยโรค")
    body(doc, "โมดูลนี้เป็นโมดูลเดียวในระบบที่รับผลของโมดูลอื่นเข้ามาเป็นตัวแปรคำนวณ "
              "โดยรับรายการรอยโรคจากตัวตรวจจับ แล้วตรวจว่ามีรอยโรคชนิดการเปลี่ยนแปลงไขมันเฉพาะที่ "
              "หรือชนิดการเว้นไขมันเฉพาะที่หรือไม่ หากพบจะบวกคะแนน 1.8")
    lead(doc, "เหตุผลทางคลินิก: ",
         "รอยโรคทั้งสองชนิดเป็นหลักฐานโดยตรงว่ามีการสะสมไขมันในตับ แม้ค่าความสว่างเฉลี่ยทั้งก้อน "
         "อาจยังไม่ถึงเกณฑ์ เนื่องจากการสะสมกระจุกอยู่เฉพาะบางบริเวณจึงถูกเฉลี่ยจนเจือจาง "
         "คะแนน 1.8 ที่ให้จึงเพียงพอที่จะยกระดับจาก S0 ไปถึง S2 ได้ด้วยตัวเอง")
    lead(doc, "ผลต่อลำดับการทำงาน: ",
         "การพึ่งพานี้บังคับให้โมดูลไขมันต้องรอผลของโมดูลรอยโรคก่อน จึงเป็นจุดพึ่งพาข้ามโมดูลเพียงจุดเดียว "
         "ในระบบทั้งหมด ส่วนโมดูลอื่นทำงานขนานกันได้อย่างอิสระ รายละเอียดอยู่ในหัวข้อ 11.5")

    h2(doc, "5.7 การแก้ปัญหาการประเมินเกินจริงในภาพตับปกติ")
    body(doc, "การปรับปรุงล่าสุดของโมดูลนี้แก้ที่ต้นเหตุสองทาง คือแก้บริเวณที่สุ่มตัวอย่าง และยกระดับ "
              "เกณฑ์การให้คะแนนทุกชั้น")
    table(
        doc,
        ["ประเด็น", "ค่าเดิม", "ค่าปัจจุบัน"],
        [
            ("เกณฑ์ความสูงตับขั้นต่ำ", "มากกว่า 20 พิกเซล", "มากกว่า 30 พิกเซล"),
            ("แถบตื้น", "0 % ถึง 35 % (รวมยอดกรวย)", "15 % ถึง 45 %"),
            ("แถบลึก", "65 % ถึง 100 % (รวมแถบสัญญาณตก)", "55 % ถึง 90 %"),
            ("ชั้นคะแนนอัตราการลดทอน", "1.35 / 1.20 / 1.10", "1.45 / 1.30 / 1.18"),
            ("ชั้นคะแนนความสว่าง", "110 / 85 / 68", "120 / 95 / 80"),
            ("คะแนนชั้นความสว่างต่ำสุด", "+0.4", "+0.3"),
            ("คะแนนไขมันเฉพาะที่", "+1.2", "+1.8"),
            ("เกณฑ์แบ่งระดับ S1 / S2 / S3", "0.8 / 1.6 / 2.6", "ไม่เปลี่ยนแปลง"),
        ],
        widths=[2.2, 2.2, 2.1],
    )
    body(doc, "กลไกของการแก้ไขมีสองส่วน ส่วนแรกคือการตัดยอดกรวยสัญญาณออกจากแถบตื้น "
              "ทำให้ค่าความสว่างของแถบตื้นไม่ถูกดันสูงเกินจริงด้วยแถบสว่างใกล้หัวตรวจ "
              "และการตัดแถบล่างสุดออกจากแถบลึก ทำให้ค่าความสว่างของแถบลึกไม่ถูกกดต่ำเกินจริง "
              "ด้วยบริเวณที่สัญญาณหายไป ทั้งสองอย่างเคยทำให้อัตราการลดทอนของตับปกติสูงเกินความจริง")
    lead(doc, "ตัวอย่างเชิงตัวเลข: ",
         "ภาพตับปกติที่วัดได้ค่าอัตราการลดทอน 1.25 และค่าความสว่าง 90 "
         "ด้วยเกณฑ์เดิมจะได้คะแนน 1.5 บวก 0.9 เท่ากับ 2.4 ซึ่งจัดเป็นระดับ S2 "
         "แต่ด้วยเกณฑ์ปัจจุบันจะได้คะแนน 0.0 และจัดเป็นระดับ S0 อย่างถูกต้อง "
         "ส่วนการยกคะแนนไขมันเฉพาะที่จาก 1.2 เป็น 1.8 เป็นการชดเชยให้กรณีที่มีหลักฐานตรงยังคง "
         "ถูกรายงานตามความรุนแรงที่ควรเป็น")

    h2(doc, "5.8 แผนภาพที่ 4 การแบ่งแถบความลึกและการคำนวณ")
    dbox(doc, ["หน้ากากตับที่หักถุงน้ำดีออกแล้ว",
               "หาช่วงความสูงของตับจากพิกเซลบนสุดถึงล่างสุด"], width=5.4)
    darrow(doc)
    drow(
        doc,
        [
            ["แถบตื้น  15 % – 45 %", "ตัดทับด้วยหน้ากากตับ", "ตัดปลาย 10 % และ 90 %",
             "ได้ค่าเฉลี่ยแถบตื้น"],
            ["แถบลึก  55 % – 90 %", "ตัดทับด้วยหน้ากากตับ", "ตัดปลาย 10 % และ 90 %",
             "ได้ค่าเฉลี่ยแถบลึก"],
        ],
        widths=[2.7, 2.7],
    )
    darrow(doc)
    dbox(doc, ["ตัวชี้วัด",
               "A = ค่าเฉลี่ยแถบตื้น ÷ ค่าเฉลี่ยแถบลึก",
               "E = ค่าเฉลี่ยตัดปลายของทั้งหน้ากากตับ"], width=5.4)
    darrow(doc)
    drow(
        doc,
        [
            ["จากอัตราการลดทอน", "+2.2 / +1.5 / +0.8"],
            ["จากความสว่าง", "+1.6 / +0.9 / +0.3"],
            ["จากไขมันเฉพาะที่", "+1.8"],
        ],
        widths=[2.2, 2.0, 2.3],
    )
    darrow(doc, "รวมคะแนน")
    dbox(doc, ["แบ่งระดับ", "≥ 2.6 → S3    ·    ≥ 1.6 → S2    ·    ≥ 0.8 → S1    ·    ต่ำกว่านั้น → S0"],
         width=5.4, heavy=True)
    caption(doc, "แผนภาพที่ 4 แถบทั้งสองนิยามเป็นสัดส่วนของหน้ากากตับ จึงไม่ผันแปรตามการตั้งค่าเครื่องตรวจ")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 6
# ---------------------------------------------------------------------------

def section_6(doc):
    h1(doc, "6. จุดร่วมภายในภาค ก")
    body(doc, "แม้โมดูลพังผืดจะเป็นโครงข่ายประสาทเทียมที่ฝึกด้วยข้อมูล ส่วนโมดูลไขมันเป็นการคำนวณ "
              "เชิงฟิสิกส์ แต่ทั้งสองถูกออกแบบให้ใช้ฐานร่วมกันหลายจุด ดังนี้")
    table(
        doc,
        ["จุดร่วม", "โมดูลพังผืด", "โมดูลไขมันพอกตับ"],
        [
            ("แหล่งของหน้ากากตับ", "หน้ากากชุดเดียวกันจากหัวข้อ 3.6", "หน้ากากชุดเดียวกันจากหัวข้อ 3.6"),
            ("การหักถุงน้ำดี", "ได้รับหน้ากากที่หักแล้ว", "ได้รับหน้ากากที่หักแล้ว"),
            ("ด่านบังคับก่อนทำงาน", "ผ่านด่านอวัยวะและเกณฑ์พื้นที่ตับ", "ผ่านด่านอวัยวะและเกณฑ์พื้นที่ตับ"),
            ("การจัดการหน้ากากหลายชิ้น", "ปิดช่องว่าง เก็บชิ้นใหญ่สุด เติมรู\n(ตัวประกอบขนาด 15)",
             "ใช้หน้ากากตามที่ได้รับ\nอาศัยการตัดปลายควอนไทล์แทน"),
            ("ขอบเขตเชิงพื้นที่", "กรอบสี่เหลี่ยมของหน้ากาก ขยาย 5 %", "กรอบสี่เหลี่ยมของหน้ากาก\nแล้วแบ่งแถบความลึก"),
            ("การจัดการพิกเซลนอกตับ", "กำหนดเป็นศูนย์ก่อนครอบตัด", "ไม่ถูกเลือกเข้ามาคำนวณตั้งแต่ต้น"),
            ("รูปแบบผลลัพธ์", "โครงสร้างพื้นที่ชุดเดียวกัน\nรหัสนำหน้า fib-", "โครงสร้างพื้นที่ชุดเดียวกัน\nรหัสนำหน้า stea-"),
            ("การรับมือกรณีข้อมูลบกพร่อง", "หน้ากากว่าง → คืนระดับ F0", "หน้ากากว่าง → คืนระดับ S0"),
        ],
        widths=[1.6, 2.5, 2.4],
    )
    lead(doc, "ข้อสังเกตเชิงเทคนิคที่ต่างกัน: ",
         "โมดูลพังผืดสร้างภาพขึ้นมาใหม่จริง จึงมีพิกเซลสีดำของบริเวณนอกตับอยู่ในเทนเซอร์ที่ป้อนเข้าโครงข่าย "
         "ทำให้โครงข่ายมองเห็นรูปทรงของตับไปพร้อมกับพื้นผิว ซึ่งเป็นเหตุผลที่ต้องมีชุดควบคุมเชิงลบ "
         "ที่ใช้เฉพาะรูปทรงหน้ากากตามหัวข้อ 4.12 ส่วนโมดูลไขมันไม่เคยประกอบภาพขึ้นมาเลย "
         "แต่คำนวณค่าสถิติจากพิกเซลที่ถูกเลือกโดยตรง จึงไม่มีช่องทางให้รูปทรงรั่วเข้าไปในผลลัพธ์")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 7 -- focal lesion detector
# ---------------------------------------------------------------------------

def section_7(doc):
    h1(doc, "7. ตัวตรวจจับรอยโรคเฉพาะที่ในเนื้อตับ")
    body(doc, "โมดูลนี้เป็นโครงข่ายตรวจจับวัตถุที่ฝึกขึ้นใหม่จากภาพอัลตราซาวด์ตับ ทำหน้าที่ระบุทั้งตำแหน่ง "
              "และชนิดของรอยโรคในเนื้อตับ จึงเป็นโมดูลที่ตอบคำถามเรื่องการระบุพื้นที่ของรอยโรคโดยตรง "
              "สิ่งที่แตกต่างจากการนำตัวตรวจจับวัตถุทั่วไปมาใช้ อยู่ที่การปรับการเพิ่มความหลากหลายของข้อมูล "
              "ให้ตรงกับฟิสิกส์ของภาพอัลตราซาวด์ และกลไกกักบริเวณเชิงพื้นที่ที่อธิบายในหัวข้อ 7.4")

    h2(doc, "7.1 โครงสร้างและการตั้งค่า")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("สถาปัตยกรรม", "YOLOv8 รุ่นเล็กที่สุด หัวทำนายแบบแยกส่วนและไม่ใช้กล่องยึด"),
            ("น้ำหนักตั้งต้น", "น้ำหนักที่ฝึกมาก่อนจากชุดภาพทั่วไป"),
            ("ขนาดภาพขณะฝึก", "512 × 512 พิกเซล"),
            ("ขนาดภาพขณะอนุมาน", "640 × 640 พิกเซล"),
            ("จำนวนคลาส", "7 ชนิด"),
            ("ขนาดไฟล์น้ำหนัก", "ประมาณ 6.2 เมกะไบต์"),
            ("ค่าความมั่นใจขั้นต่ำ", "0.25 (ผู้เรียกกำหนดค่าอื่นได้)"),
            ("เกณฑ์ IoU ของการกดทับกล่องซ้อน", "0.7 ทำแยกรายคลาส"),
            ("จำนวนกล่องสูงสุดต่อภาพ", "300"),
            ("การทำนายซ้ำหลายมุมมอง", "ไม่ใช้"),
        ],
        widths=[2.4, 4.1],
    )
    lead(doc, "หมายเหตุ: ",
         "ระบบใช้ค่าความมั่นใจขั้นต่ำค่าเดียวกับทุกคลาส ไม่มีการกำหนดเกณฑ์แยกรายชนิดรอยโรค "
         "และขนาดภาพขณะอนุมานต่างจากขณะฝึก ซึ่งเป็นประเด็นที่ระบุไว้ในภาคผนวก ค")

    h2(doc, "7.2 ระบบชื่อคลาสรอยโรค")
    body(doc, "ระบบนิยามคลาสรอยโรคเชิงคลินิกไว้เจ็ดชนิด ตามที่ใช้ในชั้นแปลผลและชั้นกฎรวมผล ดังนี้")
    table(
        doc,
        ["ลำดับ", "ชนิดรอยโรค", "ความหมาย"],
        [
            ("0", "FFC", "การเปลี่ยนแปลงไขมันเฉพาะที่"),
            ("1", "FFS", "การเว้นไขมันเฉพาะที่"),
            ("2", "HCC", "มะเร็งตับชนิดปฐมภูมิ"),
            ("3", "Cyst", "ถุงน้ำ"),
            ("4", "Hemangioma", "เนื้องอกหลอดเลือดชนิดไม่ร้าย"),
            ("5", "Dysplastic", "ก้อนเนื้อที่มีการเปลี่ยนแปลงของเซลล์"),
            ("6", "CCA", "มะเร็งท่อน้ำดีในตับ"),
        ],
        widths=[0.8, 2.0, 3.7],
        align_cols=[WD_ALIGN_PARAGRAPH.CENTER, None, None],
    )
    lead(doc, "ประเด็นที่ต้องยืนยันก่อนยกร่างข้อถือสิทธิ: ",
         "ไฟล์น้ำหนักที่ใช้งานอยู่ในปัจจุบันบันทึกรายชื่อคลาสไว้เป็นอีกชุดหนึ่ง คือ Hemangioma, Cyst, "
         "Calcification, Metastasis, HCC, CCA และ FFC ซึ่งไม่ตรงกับตารางข้างต้นทั้งลำดับและรายการ "
         "ระบบจึงมีชั้นแปลงชื่อคลาสคั่นอยู่ และการแปลงนั้นยังไม่สอดคล้องกัน "
         "เอกสารฉบับนี้บรรยายระบบชื่อคลาสที่ตั้งใจออกแบบไว้ และระบุชั้นแปลงเป็นรายละเอียดการนำไปใช้ "
         "รายละเอียดของประเด็นนี้อยู่ในภาคผนวก ค ข้อ 1")

    h2(doc, "7.3 ชุดข้อมูลและขั้นตอนการฝึก")
    table(
        doc,
        ["รายการ", "ค่า", "รายการ", "ค่า"],
        [
            ("จำนวนภาพในชุดต้นทาง", "14,448", "จำนวนรอบการฝึก", "35"),
            ("จำนวนไฟล์ป้ายกำกับ", "7,222", "ขนาดชุดย่อย", "32"),
            ("จำนวนกล่องขอบเขต", "8,574", "รอบอดทนก่อนหยุด", "15"),
            ("สัดส่วนแบ่งฝึกและตรวจสอบ", "85 : 15", "อัตราเรียนรู้เริ่มต้นและสุดท้าย", "0.01 และ 0.01"),
            ("ค่าเมล็ดสุ่ม", "42", "โมเมนตัม", "0.937"),
            ("อุปกรณ์ประมวลผล", "หน่วยเร่งบนชิป Apple", "ค่าลดทอนน้ำหนัก", "0.0005"),
            ("รอบอุ่นเครื่อง", "3.0", "น้ำหนักสูญเสีย กล่อง/คลาส/DFL", "7.5 / 0.5 / 1.5"),
            ("ปิด mosaic ช่วงท้าย", "10 รอบสุดท้าย", "เวลาที่ใช้ฝึกทั้งสิ้น", "ประมาณ 12,154 วินาที"),
        ],
        widths=[1.85, 1.5, 1.85, 1.3],
    )
    body(doc, "การกระจายจำนวนกล่องตามชนิดรอยโรคในชุดที่ใช้ฝึกจริง มีดังนี้ Cyst 1,786 กล่อง "
              "Calcification 1,712 กล่อง Metastasis 1,329 กล่อง HCC 1,311 กล่อง Hemangioma 946 กล่อง "
              "FFC 873 กล่อง และ CCA 617 กล่อง")
    lead(doc, "ข้อจำกัดของการแบ่งข้อมูล: ",
         "การแบ่งชุดฝึกและชุดตรวจสอบของโมดูลนี้ใช้การสุ่มสลับลำดับไฟล์แบบแบน ไม่ได้แบ่งตามผู้ป่วย "
         "ต่างจากโมดูลพังผืดที่แบ่งตามรหัสประจำตัวและมีการตรวจการรั่วไหลอย่างเข้มงวด "
         "จึงไม่ควรกล่าวอ้างในคำขอว่าโมดูลนี้ปลอดการรั่วไหลระดับผู้ป่วย รายละเอียดอยู่ในภาคผนวก ค ข้อ 3")
    page_break(doc)

    h2(doc, "7.4 การเพิ่มความหลากหลายของข้อมูลที่ปรับเฉพาะภาพอัลตราซาวด์")
    body(doc, "ค่าการเพิ่มความหลากหลายของข้อมูลถูกปรับให้ต่างจากค่าตั้งต้นของไลบรารีอย่างมีเหตุผลรองรับ "
              "เนื่องจากค่าตั้งต้นนั้นออกแบบมาสำหรับภาพถ่ายทั่วไป ไม่ใช่ภาพอัลตราซาวด์")
    table(
        doc,
        ["พารามิเตอร์", "ค่าที่ใช้", "เหตุผลเชิงเทคนิค"],
        [
            ("พลิกภาพบนล่าง", "0.0 คือปิดสนิท",
             "แกนตั้งของภาพอัลตราซาวด์คือทิศทางความลึกของลำคลื่น\n"
             "การพลิกทำให้ลำดับการลดทอนสัญญาณกลับหัว\nซึ่งเป็นภาพที่เป็นไปไม่ได้ทางกายภาพ"),
            ("พลิกภาพซ้ายขวา", "0.5",
             "ทิศทางซ้ายขวาขึ้นกับการวางหัวตรวจ ไม่มีความหมายคงที่"),
            ("mosaic", "0.5 ลดจากค่าตั้งต้น 1.0",
             "การต่อภาพสี่ภาพเข้าด้วยกันทำลายบริบททางกายวิภาค\nจึงลดสัดส่วนลงแทนการปิดทั้งหมด"),
            ("ความสว่างของภาพ", "0.3",
             "จำลองการปรับค่าขยายสัญญาณของเครื่องตรวจ\nซึ่งเป็นความแปรผันที่พบจริงระหว่างเครื่อง"),
            ("ความอิ่มสีและเฉดสี", "0.2 และ 0.015",
             "ตั้งค่าไว้ต่ำ เนื่องจากภาพ B-mode เป็นภาพระดับสีเทา"),
            ("หมุนภาพ", "10 องศา", "จำลองการเอียงหัวตรวจตามการใช้งานจริง"),
            ("เลื่อนและย่อขยาย", "0.10 และ 0.15", "จำลองตำแหน่งและระยะวางหัวตรวจที่ต่างกันเล็กน้อย"),
            ("mixup", "0.1", "ใช้เพียงเล็กน้อยเพื่อไม่ให้ลักษณะรอยโรคเลือนหาย"),
        ],
        widths=[1.5, 1.55, 3.45],
    )

    h2(doc, "7.5 กลไกกักบริเวณเชิงพื้นที่")
    body(doc, "หลังการกดทับกล่องซ้อนแล้ว ระบบยังไม่รับกล่องที่ได้ทันที แต่ส่งผ่านตัวกรองสองชั้นที่อ้างอิง "
              "หน้ากากตับ กลไกนี้เป็นส่วนที่ออกแบบขึ้นเฉพาะสำหรับงานนี้ และเป็นสาระสำคัญของโมดูล")
    numbered(doc, 1, "ทิ้งกล่องที่มีความกว้างหรือความสูงน้อยกว่า 6 พิกเซล เนื่องจากขนาดระดับนั้น "
                     "เป็นจุดรบกวนแบบ speckle ไม่ใช่รอยโรคที่มีความหมายทางคลินิก",
                     label="ชั้นที่ 1 กรองจุดรบกวน: ")
    numbered(doc, 2, "เก็บกล่องไว้ก็ต่อเมื่อ จุดกึ่งกลางของกล่องตกอยู่บนพิกเซลที่เป็นเนื้อตับ "
                     "หรือ สัดส่วนพื้นที่ของกล่องที่ทับกับหน้ากากตับไม่น้อยกว่า 0.25 "
                     "หากไม่เข้าเงื่อนไขใดเลย ให้ทิ้งกล่องนั้น",
                     label="ชั้นที่ 2 กักบริเวณด้วยหน้ากากตับ: ")
    body(doc, "เงื่อนไขสองข้อในชั้นที่สองเชื่อมด้วยหรือ โดยตั้งใจให้จุดกึ่งกลางเป็นทางผ่านสำรอง "
              "เพื่อรองรับรอยโรคที่อยู่ชิดขอบตับ ซึ่งกล่องขอบเขตย่อมล้ำออกนอกหน้ากากไปมาก "
              "จนสัดส่วนการทับซ้อนต่ำกว่าเกณฑ์ แต่ตัวรอยโรคยังอยู่ในเนื้อตับจริง")
    lead(doc, "ผลที่ได้: ",
         "กล่องที่เกิดจากซี่โครง เงาอะคูสติก ลมในลำไส้ หรือผนังหน้าท้อง ถูกตัดออกตั้งแต่ชั้นนี้ "
         "และเนื่องจากหน้ากากตับถูกหักพิกเซลถุงน้ำดีออกไปแล้วสองรอบตามหัวข้อ 3.6 "
         "สิ่งที่อยู่ภายในถุงน้ำดี เช่น นิ่ว จึงไม่ถูกรายงานว่าเป็นรอยโรคในเนื้อตับ "
         "ชุดทดสอบของโครงการยืนยันคุณสมบัตินี้ด้วยการป้อนหน้ากากตับว่างเปล่าพร้อมลดค่าความมั่นใจ "
         "ขั้นต่ำลงเหลือ 0.10 แล้วตรวจว่าระบบต้องไม่รายงานรอยโรคใดเลย")

    h2(doc, "7.6 การประมาณขนาดและการรายงานตำแหน่ง")
    bullet(doc, "คำนวณจากด้านที่ยาวกว่าของกล่อง คูณด้วยค่าคงที่ 0.35 มิลลิเมตรต่อพิกเซล "
                "ค่าคงที่นี้จำเป็นเพราะระบบไม่รองรับ DICOM จึงไม่มีค่าระยะห่างต่อพิกเซลจากเครื่องตรวจ "
                "ขนาดที่รายงานจึงเป็นค่าประมาณ ไม่ใช่ค่าวัด", label="ขนาดรอยโรค: ")
    bullet(doc, "รายงานเป็นพิกัดปรับมาตรฐานสองคู่ คือมุมบนซ้ายและมุมล่างขวา "
                "ค่าทุกค่าถูกจำกัดให้อยู่ในช่วง 0 ถึง 1 ระบบไม่แปลงตำแหน่งเป็นเซกเมนต์ตับตามระบบ "
                "Couinaud เนื่องจากไม่มีการตรวจหาจุดสังเกตทางกายวิภาคที่จำเป็นต่อการแบ่งเซกเมนต์",
           label="ตำแหน่ง: ")
    bullet(doc, "กล่องที่มีค่าความมั่นใจต่ำกว่า 0.60 จะถูกติดธงและทำให้ระบบออกคำเตือน "
                "แต่ยังคงถูกรายงานต่อผู้ใช้ ธงนี้จึงเป็นการแจ้งเตือน ไม่ใช่ตัวกรอง",
           label="ธงค่าความมั่นใจต่ำ: ")
    bullet(doc, "หากตัวตรวจจับเกิดข้อผิดพลาด ระบบจะคืนผลว่าไม่พบรอยโรคพร้อมติดธงความมั่นใจต่ำ "
                "และแนบข้อความข้อผิดพลาดไว้ในเหตุผลประกอบ", label="พฤติกรรมเมื่อเกิดข้อผิดพลาด: ")

    h2(doc, "7.7 ผลการฝึกที่วัดได้")
    table(
        doc,
        ["ตัวชี้วัด", "ค่าที่รอบสุดท้าย", "หมายเหตุ"],
        [
            ("ความแม่นยำ", "0.767", "สัดส่วนกล่องที่รายงานแล้วถูกต้อง"),
            ("ความไว", "0.611", "สัดส่วนรอยโรคจริงที่ตรวจพบ"),
            ("mAP ที่ IoU 0.50", "0.681", "ค่าสูงสุดระหว่างการฝึกคือ 0.691 ที่รอบที่ 31"),
            ("mAP ที่ IoU 0.50 ถึง 0.95", "0.336", "สะท้อนความแม่นของตำแหน่งกล่อง"),
        ],
        widths=[1.9, 1.5, 3.1],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    body(doc, "ค่าความไวที่ 0.611 สะท้อนว่าโมดูลนี้เหมาะกับบทบาทเครื่องมือช่วยชี้จุดให้แพทย์ตรวจซ้ำ "
              "ไม่ใช่เครื่องมือคัดกรองที่ใช้ตัดสินว่าไม่มีรอยโรคได้ด้วยตัวเอง "
              "ระบบจึงกำหนดให้ผลของโมดูลนี้ต้องผ่านการตรวจทานของแพทย์เสมอ")
    page_break(doc)

    h2(doc, "7.8 แผนภาพที่ 5 สายงานตัวตรวจจับรอยโรคและกลไกกักบริเวณ")
    dbox(doc, ["ข้อมูลเข้า", "ภาพสีที่ผ่านการลบข้อมูลระบุตัวตน  +  หน้ากากตับที่หักถุงน้ำดีแล้ว"],
         width=5.6)
    darrow(doc)
    dbox(doc, ["ตัวตรวจจับ YOLOv8", "อนุมานที่ขนาด 640 × 640 · ค่าความมั่นใจขั้นต่ำ 0.25",
               "กดทับกล่องซ้อนแยกรายคลาสที่ IoU 0.7 · สูงสุด 300 กล่อง"], width=5.6)
    darrow(doc, "รายการกล่องดิบ")
    dbox(doc, ["ชั้นที่ 1  กรองจุดรบกวน",
               "ทิ้งกล่องที่ความกว้างหรือความสูงน้อยกว่า 6 พิกเซล"], width=5.6)
    darrow(doc)
    dbox(doc, ["ชั้นที่ 2  กักบริเวณด้วยหน้ากากตับ",
               "เก็บไว้เมื่อ  จุดกึ่งกลางอยู่บนเนื้อตับ",
               "หรือ  สัดส่วนพื้นที่กล่องที่ทับหน้ากากตับ ≥ 0.25",
               "มิฉะนั้นทิ้ง (ซี่โครง · เงาอะคูสติก · ลมในลำไส้ · ผนังหน้าท้อง)"],
         width=5.6, heavy=True)
    darrow(doc)
    dbox(doc, ["ผลลัพธ์ต่อกล่องที่ผ่าน",
               "ชนิดรอยโรค · ค่าความมั่นใจ · พิกัดปรับมาตรฐาน 0–1",
               "ขนาดโดยประมาณ = ด้านที่ยาวกว่า × 0.35 มิลลิเมตรต่อพิกเซล",
               "ธงเตือนเมื่อค่าความมั่นใจต่ำกว่า 0.60"], width=5.6)
    darrow(doc)
    drow(
        doc,
        [
            ["ส่งต่อโมดูลไขมัน", "เมื่อพบ FFC หรือ FFS", "บวกคะแนน 1.8"],
            ["ส่งต่อชั้นกฎรวมผล", "เมื่อพบ HCC หรือ CCA", "ยกระดับตามหัวข้อ 8"],
        ],
        widths=[2.7, 2.9],
    )
    caption(doc, "แผนภาพที่ 5 กล่องเส้นหนาคือกลไกกักบริเวณเชิงพื้นที่ ซึ่งเป็นชั้นที่ตัดผลบวกลวง "
                 "จากโครงสร้างนอกตับออกทั้งหมด")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 8 -- malignancy escalation
# ---------------------------------------------------------------------------

def section_8(doc):
    h1(doc, "8. กฎยกระดับเมื่อพบรอยโรคมะเร็ง และการส่งตรวจยืนยัน")
    body(doc, "หัวข้อนี้อธิบายกลไกที่ผลของโมดูลรอยโรคย้อนกลับไปแก้ไขผลของโมดูลพังผืด "
              "ซึ่งเป็นจุดเดียวในระบบที่ผลของโมดูลหนึ่งเขียนทับผลของอีกโมดูลหนึ่ง "
              "และเป็นกฎที่ออกแบบขึ้นใหม่โดยอ้างอิงความสัมพันธ์ทางคลินิกระหว่างมะเร็งตับกับภาวะตับแข็ง")

    h2(doc, "8.1 เหตุผลทางคลินิก")
    body(doc, "มะเร็งตับชนิดปฐมภูมิเกิดขึ้นบนพื้นหลังของตับแข็งเป็นส่วนใหญ่ หากระบบตรวจพบรอยโรค "
              "ที่เข้าได้กับมะเร็งชนิดนี้ แต่โมดูลพังผืดกลับรายงานว่าเนื้อตับอยู่ในระดับต้น "
              "ความไม่สอดคล้องนั้นมีแนวโน้มเกิดจากข้อจำกัดของโมดูลพังผืดที่หดค่าทำนายเข้าหาค่ากลาง "
              "ตามที่อธิบายไว้ในหัวข้อ 4.13 มากกว่าจะเป็นเพราะเนื้อตับปกติจริง")
    body(doc, "ระบบจึงเลือกยกระดับพื้นหลังของเนื้อตับขึ้น เพื่อไม่ให้ผลรายงานสื่อไปในทางที่ปลอดภัย "
              "เกินความเป็นจริง ซึ่งเป็นทิศทางตรงข้ามกับการแก้ปัญหาการประเมินเกินจริงในหัวข้อ 4.11 "
              "และเป็นความตั้งใจในการออกแบบ กล่าวคือ ระบบเข้มงวดขึ้นกับภาพปกติ "
              "และระมัดระวังมากขึ้นเมื่อมีหลักฐานของรอยโรคร้าย")

    h2(doc, "8.2 เงื่อนไขและผลของกฎ")
    table(
        doc,
        ["ขั้น", "รายละเอียด"],
        [
            ("เงื่อนไขกระตุ้น",
             "พบรอยโรคชนิด HCC หรือชนิด CCA อย่างน้อยหนึ่งจุดในผลของโมดูลรอยโรค\n"
             "โดยไม่มีการกำหนดค่าความมั่นใจขั้นต่ำเพิ่มเติมสำหรับกฎนี้"),
            ("เงื่อนไขเพิ่มเติมของโมดูลพังผืด",
             "ระดับความเสี่ยงต่ำกว่าระดับสูง  หรือ  ระดับที่รายงานอยู่ใน F0, F1 หรือ F2\n"
             "หากอยู่ที่ระดับ F3 พร้อมความเสี่ยงสูงอยู่แล้ว กฎนี้จะไม่แก้ไขค่าใด"),
            ("การแก้ไขที่เกิดขึ้น",
             "ระดับ → F4 (เขียนทับ)\n"
             "ระดับความเสี่ยง → สูง พร้อมข้อความกำกับว่าเป็นความเสี่ยงจากพื้นหลังตับแข็ง\n"
             "ค่าประมาณ kPa → ค่าที่มากกว่าระหว่างค่าเดิมกับ 9.5 (ยกพื้น ไม่เขียนทับ)\n"
             "โอกาสเป็น F4 → ค่าที่มากกว่าระหว่างค่าเดิมกับ 0.65 (ยกพื้น ไม่เขียนทับ)"),
            ("คำเตือนที่ออก",
             "แนะนำส่งตรวจยืนยันด้วย CT Triphasic Liver Protocol หรือ MRI ตับ\n"
             "และตรวจระดับสารบ่งชี้มะเร็งในซีรัม คือ AFP และ CA19-9 โดยด่วน"),
        ],
        widths=[1.8, 4.7],
    )
    body(doc, "การใช้การยกพื้นแทนการเขียนทับสำหรับค่า kPa และค่าโอกาสเป็น F4 เป็นการออกแบบที่ตั้งใจ "
              "เพื่อไม่ให้กฎนี้ลดค่าที่โมเดลทำนายไว้สูงกว่าเกณฑ์อยู่แล้วให้ต่ำลง "
              "กฎนี้จึงยกระดับได้อย่างเดียว ไม่มีทางลดระดับ")

    h2(doc, "8.3 ลำดับการทำงานที่เกี่ยวข้อง")
    body(doc, "กฎนี้ทำงานในชั้นตรวจทานหลังจากโมดูลทั้งสี่ทำงานเสร็จแล้ว และทำงานก่อนชั้นเรียบเรียงรายงาน "
              "จึงมีผลว่ารายงานข้อความที่ผู้ใช้ได้รับสะท้อนค่าที่ผ่านการยกระดับแล้ว ไม่ใช่ค่าดิบจากโมเดล "
              "ลำดับนี้จำเป็น เพราะหากสลับกันจะเกิดกรณีที่รายงานข้อความบอกระดับหนึ่ง "
              "แต่ค่าตัวเลขในผลลัพธ์เป็นอีกระดับหนึ่ง")

    h2(doc, "8.4 คำเตือนอื่นที่ชั้นกฎรวมผลสร้างขึ้น")
    table(
        doc,
        ["เงื่อนไข", "คำเตือนที่ออก"],
        [
            ("มีรอยโรคที่ค่าความมั่นใจต่ำกว่า 0.60", "แจ้งว่าผลมีความมั่นใจต่ำ ควรให้แพทย์ตรวจทานซ้ำ"),
            ("สัดส่วนพื้นที่ตับมากกว่า 0 แต่ต่ำกว่า 5 %", "แจ้งว่าพื้นที่เนื้อตับในภาพน้อย ผลอาจไม่น่าเชื่อถือ"),
            ("ชั้นตรวจคุณภาพภาพไม่ผ่าน", "แจ้งว่าคุณภาพภาพต่ำกว่าเกณฑ์"),
            ("ระดับไขมันพอกตับเป็น S3", "แจ้งว่าการลดทอนสัญญาณในระดับลึกอาจบดบังรอยโรค\n"
                                        "และแนะนำติดตามด้วยการวัดค่า CAP หรือ FibroScan"),
        ],
        widths=[2.5, 4.0],
    )
    body(doc, "คำเตือนทั้งหมดถูกนำมาต่อกันเป็นข้อความเดียวในเขตข้อมูลคำเตือนทางคลินิกของผลลัพธ์")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 9 -- fluke risk
# ---------------------------------------------------------------------------

def section_9(doc):
    h1(doc, "9. โมดูลประเมินความเสี่ยงพยาธิใบไม้ตับและมะเร็งท่อน้ำดี")

    h2(doc, "9.1 ลักษณะทางเทคนิคของโมดูล")
    lead(doc, "ข้อเท็จจริงที่ต้องระบุให้ชัด: ",
         "โมดูลนี้ไม่ใช่โมเดลการเรียนรู้ และไม่ได้วิเคราะห์ลักษณะของภาพ แต่เป็นกฎให้คะแนนเชิงเส้น "
         "ที่คำนวณจากประวัติผู้ป่วยล้วน ค่าคะแนนที่ได้จึงขึ้นกับข้อมูลประวัติที่ป้อนเข้ามาเท่านั้น "
         "การเปลี่ยนภาพโดยที่ประวัติเหมือนเดิม ไม่ทำให้ผลของโมดูลนี้เปลี่ยน")
    body(doc, "เอกสารประกอบเดิมของโครงการบางฉบับบรรยายว่าโมดูลนี้วิเคราะห์ลักษณะพื้นผิวรอบท่อน้ำดี "
              "และคำนวณดัชนีการหนาตัวของผนังท่อ ซึ่งไม่ตรงกับสิ่งที่มีอยู่ในซอร์สโค้ดที่ใช้งานจริง "
              "เอกสารฉบับนี้จึงไม่นำข้อความดังกล่าวมาใช้ และระบุประเด็นไว้ในภาคผนวก ค ข้อ 2")

    h2(doc, "9.2 สูตรคำนวณคะแนนความเสี่ยง")
    body(doc, "คะแนนเริ่มต้นจากค่าฐาน แล้วบวกเพิ่มตามปัจจัยเสี่ยงที่พบ โดยแต่ละปัจจัยบวกได้ครั้งเดียว")
    table(
        doc,
        ["องค์ประกอบ", "คะแนน", "ที่มา"],
        [
            ("ค่าฐาน", "0.10", "ความเสี่ยงพื้นฐานของประชากรทั่วไป"),
            ("ประวัติกินปลาน้ำจืดดิบหรือสุก ๆ ดิบ ๆ", "+0.55", "ปัจจัยเสี่ยงหลักของการติดพยาธิใบไม้ตับ"),
            ("ประวัติเคยรักษาการติดพยาธิใบไม้ตับ", "+0.30", "หลักฐานการสัมผัสเชื้อในอดีต"),
            ("ประวัติมะเร็งท่อน้ำดีในครอบครัว", "+0.15", "ปัจจัยเสี่ยงร่วมของมะเร็งท่อน้ำดี"),
        ],
        widths=[3.1, 1.0, 2.4],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER, None],
    )

    h2(doc, "9.3 เกณฑ์แบ่งผลลัพธ์")
    table(
        doc,
        ["ช่วงคะแนน", "ผลที่รายงาน", "ระดับความเสี่ยง", "ค่าความมั่นใจ"],
        [
            ("ตั้งแต่ 0.70", "Probable", "สูง", "0.88"),
            ("ตั้งแต่ 0.35 แต่ไม่ถึง 0.70", "Possible", "ปานกลาง", "0.78"),
            ("น้อยกว่า 0.35", "Negative", "ต่ำ", "0.92"),
        ],
        widths=[2.15, 1.55, 1.4, 1.4],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER,
                    WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )

    h2(doc, "9.4 ส่วนที่เกี่ยวข้องกับภาพ")
    body(doc, "โมดูลรับภาพระดับสีเทาและหน้ากากตับเข้ามาด้วย แต่ใช้เพียงเพื่อคำนวณจุดศูนย์ถ่วง "
              "ของหน้ากากตับ แล้ววาดเส้นอิสระสามจุดรอบจุดศูนย์ถ่วงนั้นด้วยระยะเยื้องคงที่ "
              "เพื่อทำเครื่องหมายบริเวณรอบท่อน้ำดีบนภาพซ้อนทับ "
              "เส้นดังกล่าวเป็นเครื่องหมายอ้างอิงตำแหน่งเชิงเรขาคณิต ไม่ได้มาจากการวัดค่าใดในภาพ "
              "และไม่มีผลต่อคะแนนความเสี่ยง")

    h2(doc, "9.5 ข้อจำกัดที่ต้องทราบ")
    bullet(doc, "ในช่องทางวิเคราะห์เต็มรูปแบบ โครงสร้างข้อมูลประวัติผู้ป่วยที่ระบบนิยามไว้ "
                "มีเฉพาะเขตข้อมูลอายุ เพศ ไวรัสตับอักเสบบี ไวรัสตับอักเสบซี ประวัติดื่มสุรา "
                "และการกินปลาน้ำจืดดิบ จึงไม่มีเขตข้อมูลสำหรับประวัติเคยรักษาพยาธิ "
                "และประวัติมะเร็งในครอบครัว คะแนนสูงสุดที่ทำได้ผ่านช่องทางนี้คือ 0.65 "
                "ซึ่งยังไม่ถึงเกณฑ์ 0.70 ผลจึงไม่สามารถขึ้นถึงระดับ Probable ได้")
    bullet(doc, "ระดับ Probable เข้าถึงได้เฉพาะผ่านช่องทางรายโมดูล ซึ่งรับประวัติเป็นโครงสร้างข้อมูล "
                "แบบอิสระ จึงส่งเขตข้อมูลทั้งสองที่เหลือเข้ามาได้")
    bullet(doc, "ผลของโมดูลนี้ไม่ได้ถูกนำเข้าชั้นกฎรวมผล จึงไม่ทำให้เกิดคำเตือนหรือการยกระดับใด "
                "ในฝั่งประมวลผลหลัก แต่ปรากฏในรายงานข้อความที่ส่งให้ผู้ใช้")
    body(doc, "ข้อจำกัดทั้งสามข้อนี้เป็นเรื่องของการเชื่อมต่อภายในระบบ ไม่ใช่ข้อจำกัดของหลักการ "
              "และได้ระบุไว้ในภาคผนวก ค เพื่อให้พิจารณาแก้ไขก่อนกำหนดขอบเขตข้อถือสิทธิ")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 10
# ---------------------------------------------------------------------------

def section_10(doc):
    h1(doc, "10. จุดร่วมภายในภาค ข")
    table(
        doc,
        ["จุดร่วม", "โมดูลรอยโรค", "โมดูลความเสี่ยงพยาธิ"],
        [
            ("แหล่งของหน้ากากตับ", "หน้ากากชุดเดียวกันจากหัวข้อ 3.6", "หน้ากากชุดเดียวกันจากหัวข้อ 3.6"),
            ("บทบาทของหน้ากาก", "ใช้กักบริเวณกล่องตรวจจับ", "ใช้หาจุดศูนย์ถ่วงเพื่อวางเครื่องหมาย"),
            ("ประเภททางเทคนิค", "โครงข่ายที่ฝึกด้วยข้อมูล", "กฎให้คะแนนจากประวัติผู้ป่วย"),
            ("ด่านบังคับก่อนทำงาน", "ผ่านด่านอวัยวะและเกณฑ์พื้นที่ตับ", "ผ่านด่านอวัยวะและเกณฑ์พื้นที่ตับ"),
            ("รูปแบบผลลัพธ์เชิงพื้นที่", "กล่องสองจุด พิกัดปรับมาตรฐาน\nรหัสนำหน้า les-",
             "เส้นอิสระสามจุด พิกัดปรับมาตรฐาน\nรหัสนำหน้า fluke-"),
            ("ความเชื่อมโยงกับมะเร็งท่อน้ำดี", "ตรวจจับรอยโรคชนิด CCA จากภาพ",
             "ประเมินความเสี่ยงต่อ CCA จากประวัติ"),
            ("การเข้าสู่ชั้นกฎรวมผล", "เข้า ทำให้เกิดการยกระดับตามหัวข้อ 8", "ไม่เข้า ปรากฏเฉพาะในรายงานข้อความ"),
        ],
        widths=[1.7, 2.45, 2.35],
    )
    lead(doc, "ข้อสังเกต: ",
         "ทั้งสองโมดูลในภาคนี้มุ่งไปที่มะเร็งท่อน้ำดีเหมือนกัน แต่เข้าถึงจากคนละทาง "
         "โมดูลรอยโรคมองหาก้อนที่ปรากฏในภาพแล้ว ส่วนโมดูลความเสี่ยงพยาธิประเมินโอกาสจากปัจจัยเสี่ยง "
         "ที่นำไปสู่โรคนั้น การออกแบบให้ทำงานคู่กันจึงครอบคลุมทั้งผู้ที่มีรอยโรคปรากฏแล้ว "
         "และผู้ที่ยังไม่มีรอยโรคแต่อยู่ในกลุ่มเสี่ยงสูงและควรได้รับการติดตาม")
    page_break(doc)


# ---------------------------------------------------------------------------
# Section 11 -- what the models have in common
# ---------------------------------------------------------------------------

def section_11(doc):
    h1(doc, "11. จุดร่วมของทั้งสี่โมดูล")
    body(doc, "หัวข้อนี้ตอบคำถามข้อ 3 โดยตรง แม้โมดูลทั้งสี่จะต่างกันทั้งประเภททางเทคนิคและโรคที่ตรวจ "
              "แต่ถูกออกแบบให้วางอยู่บนฐานร่วมชุดเดียวกันโดยตั้งใจ เพื่อให้ผลลัพธ์ของทุกโมดูล "
              "อ้างอิงบริเวณเดียวกัน มีรูปแบบเดียวกัน และอยู่ภายใต้เงื่อนไขความปลอดภัยชุดเดียวกัน")

    h2(doc, "11.1 หน้ากากตับชุดเดียวที่ใช้ร่วมกัน")
    body(doc, "ระบบสร้างหน้ากากตับเพียงครั้งเดียวต่อการตรวจหนึ่งครั้ง แล้วส่งอาเรย์ชุดเดียวกันนั้น "
              "ให้ทุกโมดูลใช้ ไม่มีโมดูลใดสร้างหน้ากากของตัวเองขึ้นมาใหม่ ผลที่ตามมามีสองประการ")
    bullet(doc, "ผลของทุกโมดูลอ้างอิงบริเวณเนื้อตับชุดเดียวกัน จึงนำมาวางซ้อนบนภาพเดียวกัน "
                "และเปรียบเทียบกันได้โดยตรง ไม่เกิดกรณีที่โมดูลหนึ่งเห็นตับกว้างกว่าอีกโมดูลหนึ่ง")
    bullet(doc, "การหักพิกเซลถุงน้ำดีออกทำเพียงครั้งเดียวที่ต้นทาง แล้วมีผลกับทุกโมดูลพร้อมกัน "
                "จึงไม่ต้องทำซ้ำและไม่มีโอกาสที่โมดูลใดพลาดขั้นตอนนี้")
    table(
        doc,
        ["โมดูล", "วิธีใช้หน้ากากตับ"],
        [
            ("พังผืด", "คูณภาพด้วยหน้ากาก แล้วครอบตัดตามกรอบที่ขยาย 5 % เป็น ROI ป้อนเข้าโครงข่าย"),
            ("ไขมันพอกตับ", "ใช้เป็นตัวคัดเลือกพิกเซล และใช้ช่วงความสูงของหน้ากากกำหนดแถบความลึก"),
            ("รอยโรค", "ใช้เป็นเกณฑ์กักบริเวณกล่องตรวจจับ ทั้งการทดสอบจุดกึ่งกลางและสัดส่วนการทับซ้อน"),
            ("ความเสี่ยงพยาธิ", "ใช้หาจุดศูนย์ถ่วงเพื่อวางเครื่องหมายบริเวณรอบท่อน้ำดี"),
        ],
        widths=[1.5, 5.0],
    )

    h2(doc, "11.2 ด่านบังคับร่วมก่อนทำงาน")
    body(doc, "ทุกโมดูลอยู่ใต้ด่านชุดเดียวกันตามหัวข้อ 3.5 คือด่านจำแนกอวัยวะ และเกณฑ์สัดส่วนพื้นที่ตับ "
              "ที่ต้องไม่น้อยกว่า 5.0 เปอร์เซ็นต์ ลักษณะสำคัญคือด่านนี้ถูกบังคับใช้ในระดับโครงสร้างของสายงาน "
              "หากไม่ผ่าน ระบบจะคืนผลลัพธ์ที่ระบุสถานะหยุดทำงาน โดยที่ยังไม่มีการสร้างงานประมวลผล "
              "ของโมดูลใดขึ้นเลย ไม่ใช่การสร้างขึ้นแล้วทิ้งผลภายหลัง ทุกเขตข้อมูลผลของโมดูลจะเป็นค่าว่าง "
              "และจำนวนรอยโรคเป็นศูนย์")

    h2(doc, "11.3 โครงสร้างข้อมูลผลลัพธ์ร่วม")
    body(doc, "ทุกโมดูลคืนผลด้วยซองข้อมูลรูปแบบเดียวกัน ต่างกันเพียงเขตข้อมูลเสริมเฉพาะโมดูล")
    table(
        doc,
        ["เขตข้อมูลร่วม", "ความหมาย"],
        [
            ("agentId", "รหัสโมดูลที่ให้ผล"),
            ("value", "ค่าผลลัพธ์หลัก เช่น ระดับ F, ระดับ S, รายการรอยโรค หรือระดับความเสี่ยง"),
            ("confidence", "ค่าความมั่นใจโดยรวมของโมดูล"),
            ("regions", "รายการพื้นที่ตามโครงสร้างในหัวข้อ 3.7"),
            ("rationale", "เหตุผลประกอบเป็นข้อความภาษาไทย"),
            ("modelVersion", "รุ่นของโมเดลหรือกฎที่ใช้"),
            ("inferenceMs", "เวลาที่ใช้ประมวลผล"),
            ("simulated", "ระบุว่าเป็นผลจริงหรือผลจำลอง"),
        ],
        widths=[1.6, 4.9],
    )
    body(doc, "เขตข้อมูลเสริมเฉพาะโมดูลมีดังนี้ โมดูลพังผืดเพิ่มค่าประมาณ kPa และระดับความเสี่ยง "
              "โมดูลไขมันเพิ่มอัตราการลดทอนลำคลื่น และโมดูลความเสี่ยงพยาธิเพิ่มคะแนนความเสี่ยง")
    lead(doc, "ประโยชน์เชิงโครงสร้าง: ",
         "การใช้ซองข้อมูลร่วมทำให้ส่วนแสดงผล ส่วนตรวจทานของแพทย์ และส่วนบันทึกข้อมูลย้อนกลับ "
         "เขียนขึ้นเพียงชุดเดียวและรองรับทุกโมดูล การเพิ่มโมดูลใหม่ในอนาคตจึงไม่ต้องแก้ส่วนเหล่านั้น")

    h2(doc, "11.4 ค่าคงที่และวิธีเตรียมข้อมูลที่ใช้ร่วมกัน")
    bullet(doc, "ทุกโมดูลได้รับภาพที่ผ่านการลบแถบบน 12 เปอร์เซ็นต์แล้ว", label="ภาพต้นทางเดียวกัน: ")
    bullet(doc, "โครงข่ายทุกตัวในระบบใช้ค่าเฉลี่ยและส่วนเบี่ยงเบนมาตรฐานชุดเดียวกัน "
                "ทั้งด่านจำแนกอวัยวะ โครงข่ายแบ่งส่วน และโมเดลพังผืด", label="การปรับมาตรฐานค่าพิกเซล: ")
    bullet(doc, "การหากรวยสัญญาณด้วยการตั้งเกณฑ์ความสว่าง ตามด้วยการดำเนินการทางสัณฐานวิทยา "
                "และการเก็บองค์ประกอบเชื่อมต่อที่ใหญ่ที่สุด ใช้แนวทางเดียวกันทุกจุดในระบบ",
           label="การหาขอบเขตกรวยสัญญาณ: ")
    bullet(doc, "การปิดช่องว่าง เก็บชิ้นใหญ่ที่สุด และเติมรูภายใน เป็นชุดปฏิบัติการเดียวกัน "
                "ต่างกันเพียงขนาดตัวประกอบ คือ 11 สำหรับตับ 7 สำหรับถุงน้ำดี และ 15 ในขั้นเตรียม "
                "ROI ของโมดูลพังผืด", label="การทำความสะอาดหน้ากาก: ")
    bullet(doc, "ทุกโมดูลรายงานพิกัดในช่วง 0 ถึง 1 เทียบกับขนาดเฟรม ไม่มีโมดูลใดรายงานพิกัดพิกเซลดิบ",
           label="ระบบพิกัดผลลัพธ์: ")
    page_break(doc)

    h2(doc, "11.5 ลำดับการทำงานร่วมและจุดพึ่งพา")
    body(doc, "สายงานหลักแบ่งเป็นสามขั้นที่มีลักษณะการทำงานต่างกันชัดเจน")
    numbered(doc, 1, "ตรวจด่านอวัยวะและเกณฑ์สัดส่วนพื้นที่ตับ หากไม่ผ่านจะจบการทำงานทันที",
             label="ขั้นด่านบังคับ: ")
    numbered(doc, 2, "ชั้นตรวจคุณภาพภาพ ชั้นตรวจความสมเหตุสมผลของข้อมูลคลินิก และชั้นระบุมุมตรวจ "
                     "ทำงานขนานกัน", label="ขั้นเตรียมข้อมูล: ")
    numbered(doc, 3, "โมดูลรอยโรค โมดูลพังผืด โมดูลแปลงค่าความแข็งตับที่ผู้ใช้ป้อน "
                     "และโมดูลความเสี่ยงพยาธิ ถูกส่งเข้าทำงานขนานกัน จากนั้นระบบรอผลของโมดูลรอยโรค "
                     "ก่อนจึงส่งโมดูลไขมันเข้าทำงาน เนื่องจากโมดูลไขมันต้องใช้ผลการตรวจพบไขมันเฉพาะที่ "
                     "ตามหัวข้อ 5.6 นี่คือจุดพึ่งพาข้ามโมดูลเพียงจุดเดียวในระบบทั้งหมด",
             label="ขั้นโมดูลผู้เชี่ยวชาญ: ")
    numbered(doc, 4, "ชั้นกฎรวมผล ตามด้วยชั้นเรียบเรียงรายงาน และชั้นตรวจความปลอดภัย "
                     "ทำงานเรียงตามลำดับอย่างเคร่งครัด ลำดับนี้จำเป็นเพราะชั้นกฎรวมผลสามารถแก้ไข "
                     "ค่าผลของโมดูลพังผืดได้ตามหัวข้อ 8 ชั้นเรียบเรียงรายงานจึงต้องทำงานหลังจากนั้น "
                     "เพื่อให้รายงานสะท้อนค่าที่ผ่านการแก้ไขแล้ว", label="ขั้นตรวจทาน: ")

    h2(doc, "11.6 กลไกรวมผลและกฎความปลอดภัยร่วม")
    body(doc, "ชั้นกฎรวมผลเป็นจุดบรรจบของทุกโมดูล ทำหน้าที่สามอย่างคือ รวบรวมคำเตือนจากทุกแหล่ง "
              "เข้าเป็นข้อความเดียว ยกระดับผลของโมดูลพังผืดเมื่อพบรอยโรคมะเร็ง และออกคำแนะนำ "
              "ส่งตรวจยืนยัน ลักษณะสำคัญของชั้นนี้คือเป็นกฎเชิงกำหนดล้วน ไม่มีการเรียนรู้ "
              "และไม่มีการเรียกใช้บริการภายนอก จึงให้ผลเหมือนเดิมทุกครั้งและตรวจสอบย้อนกลับได้")
    body(doc, "ชั้นเรียบเรียงรายงานทำหน้าที่แปลผลเชิงโครงสร้างเป็นข้อความภาษาไทยสำหรับแพทย์ "
              "โดยส่งเฉพาะผลลัพธ์เชิงโครงสร้างออกไป ไม่ส่งพิกเซลของภาพออกนอกระบบ "
              "และหากไม่สามารถเรียกใช้บริการภายนอกได้ ระบบมีตัวเรียบเรียงสำรองเชิงกำหนด "
              "ที่ประกอบรายงานหกหัวข้อจากผลลัพธ์เดียวกัน จึงยังคงใช้งานได้โดยไม่ต้องพึ่งบริการภายนอก")

    h2(doc, "11.7 วงจรแพทย์ตรวจทานที่ใช้ร่วมกัน")
    body(doc, "ระบบบันทึกการตรวจทานของแพทย์ด้วยโครงสร้างเดียวกันสำหรับทุกโมดูล โดยมีหลักการดังนี้")
    bullet(doc, "ผลของ AI ถูกเก็บแยกจากผลการตรวจทานของแพทย์ ระบบไม่เคยเขียนทับผลของ AI "
                "จึงเปรียบเทียบผลก่อนและหลังการแก้ไขได้เสมอ", label="ไม่เขียนทับผลเดิม: ")
    bullet(doc, "การตรวจทานแต่ละโมดูลมีสถานะสามค่า คือรอตรวจ ถูกต้อง และไม่ถูกต้อง "
                "โดยสถานะไม่ถูกต้องจะยังไม่ถือว่าเสร็จสมบูรณ์จนกว่าแพทย์จะระบุเหตุผลประกอบ",
           label="บังคับให้ระบุเหตุผล: ")
    bullet(doc, "เหตุการณ์ทุกอย่างถูกบันทึกต่อท้ายอย่างเดียว ทั้งการตั้งสถานะ การแก้ค่า "
                "การเพิ่มและลบเครื่องหมายบนภาพ การเปลี่ยนผลคัดแยก และการบันทึกหมายเหตุ",
           label="บันทึกแบบต่อท้ายอย่างเดียว: ")
    bullet(doc, "เครื่องหมายที่แพทย์วาดเพิ่มใช้โครงสร้างพื้นที่ชุดเดียวกับที่ AI ใช้ "
                "โดยระบุแหล่งที่มาเป็นแพทย์ และผูกกับโมดูลที่ต้องการแก้ไขได้",
           label="ใช้โครงสร้างพื้นที่ร่วม: ")
    body(doc, "ข้อมูลที่บันทึกไว้ประกอบด้วย ค่าที่ AI ให้ ค่าความมั่นใจ รุ่นของโมเดล ค่าที่แพทย์แก้ไข "
              "เหตุผลประกอบ จำนวนครั้งที่แก้ไข และลำดับเหตุการณ์ทั้งหมด "
              "จึงใช้เป็นชุดข้อมูลสำหรับฝึกซ้ำและสำหรับวัดความสอดคล้องระหว่างแพทย์กับ AI ได้โดยตรง")
    page_break(doc)

    h2(doc, "11.8 ตารางสรุปจุดร่วมและจุดต่างของทั้งสี่โมดูล")
    table(
        doc,
        ["ประเด็น", "พังผืด", "ไขมันพอกตับ", "รอยโรค", "ความเสี่ยงพยาธิ"],
        [
            ("ประเภททางเทคนิค", "โครงข่ายที่ฝึก\nด้วยข้อมูล", "กฎคำนวณ\nเชิงฟิสิกส์",
             "โครงข่ายที่ฝึก\nด้วยข้อมูล", "กฎจากประวัติ\nผู้ป่วย"),
            ("มีไฟล์น้ำหนัก", "มี (5 โมเดล)", "ไม่มี", "มี (1 โมเดล)", "ไม่มี"),
            ("ใช้หน้ากากตับร่วม", "ใช่", "ใช่", "ใช่", "ใช่"),
            ("อยู่ใต้ด่านบังคับ", "ใช่", "ใช่", "ใช่", "ใช่"),
            ("ใช้ซองข้อมูลร่วม", "ใช่", "ใช่", "ใช่", "ใช่"),
            ("รายงานพิกัดปรับมาตรฐาน", "ใช่ (กล่อง)", "ใช่ (กล่อง)", "ใช่ (กล่อง)", "ใช่ (เส้นอิสระ)"),
            ("ใช้ข้อมูลคลินิก", "ใช้มุมตรวจ", "ไม่ใช้", "ไม่ใช้", "ใช้ประวัติทั้งหมด"),
            ("รับผลจากโมดูลอื่น", "รับจากชั้นกฎรวมผล", "รับจากโมดูลรอยโรค", "ไม่รับ", "ไม่รับ"),
            ("เข้าชั้นกฎรวมผล", "เข้า", "เข้า", "เข้า", "ไม่เข้า"),
            ("เข้าวงจรแพทย์ตรวจทาน", "เข้า", "เข้า", "เข้า", "เข้า"),
        ],
        widths=[1.55, 1.25, 1.25, 1.25, 1.2],
        align_cols=[None] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )

    h2(doc, "11.9 แผนภาพที่ 6 ผังจุดร่วมของทั้งสี่โมดูล")
    dbox(doc, ["ฐานร่วมที่ 1  ภาพต้นทางเดียวกัน",
               "ภาพที่ผ่านการลบแถบบน 12 เปอร์เซ็นต์ · ค่าปรับมาตรฐานพิกเซลชุดเดียวกัน"], width=6.0)
    darrow(doc)
    dbox(doc, ["ฐานร่วมที่ 2  ด่านบังคับ",
               "ด่านจำแนกอวัยวะ  และ  สัดส่วนพื้นที่ตับ ≥ 5.0 เปอร์เซ็นต์",
               "ไม่ผ่าน → ไม่มีโมดูลใดถูกสร้างขึ้นทำงานเลย"], width=6.0, heavy=True)
    darrow(doc)
    dbox(doc, ["ฐานร่วมที่ 3  หน้ากากตับชุดเดียว",
               "สร้างครั้งเดียวต่อการตรวจ · หักพิกเซลถุงน้ำดีออกสองรอบ",
               "ส่งอาเรย์ชุดเดียวกันให้ทุกโมดูล"], width=6.0)
    dsplit(doc)
    drow(
        doc,
        [
            ["พังผืด", "ครอบตัด ROI", "จากหน้ากาก"],
            ["ไขมัน", "แบ่งแถบความลึก", "ในหน้ากาก"],
            ["รอยโรค", "กักบริเวณกล่อง", "ด้วยหน้ากาก"],
            ["พยาธิ", "หาจุดศูนย์ถ่วง", "ของหน้ากาก"],
        ],
        widths=[1.5, 1.5, 1.5, 1.5],
    )
    darrow(doc, "ทุกโมดูลคืนซองข้อมูลรูปแบบเดียวกัน")
    dbox(doc, ["ฐานร่วมที่ 4  ชั้นกฎรวมผลเชิงกำหนด",
               "รวมคำเตือน · ยกระดับพังผืดเมื่อพบ HCC หรือ CCA · ออกคำแนะนำส่งตรวจยืนยัน"],
         width=6.0)
    darrow(doc)
    dbox(doc, ["ฐานร่วมที่ 5  วงจรแพทย์ตรวจทาน",
               "โครงสร้างพื้นที่ร่วม · บันทึกต่อท้ายอย่างเดียว · ไม่เขียนทับผลของ AI"], width=6.0)
    caption(doc, "แผนภาพที่ 6 ฐานร่วมทั้งห้าชั้นเป็นสิ่งที่ทุกโมดูลใช้เหมือนกัน "
                 "ส่วนที่ต่างกันมีเพียงวิธีใช้หน้ากากตับในแถวกลางเท่านั้น")
    page_break(doc)


# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------

def appendix_a(doc):
    h1(doc, "ภาคผนวก ก  ตารางค่าคงที่และเกณฑ์ตัดสินทั้งหมด")
    body(doc, "รวบรวมค่าคงที่ทุกค่าที่ปรากฏในเอกสารไว้ที่เดียว เพื่อใช้อ้างอิงขณะยกร่างคำขอ")

    h2(doc, "ก.1 ด่านคัดกรองและการเตรียมข้อมูล")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("สัดส่วนแถบบนที่ลบทิ้ง", "12 % ของความสูงเฟรม"),
            ("ขนาดไฟล์อัปโหลดสูงสุด", "25 เมกะไบต์"),
            ("ค่าความมั่นใจขั้นต่ำของด่านอวัยวะ", "0.55"),
            ("ค่าเอนโทรปีสูงสุดของด่านอวัยวะ", "1.30"),
            ("สัดส่วนพื้นที่ตับขั้นต่ำ", "5.0 %"),
            ("จำนวนคลาสของด่านอวัยวะ", "10"),
            ("ขนาดภาพเข้าด่านอวัยวะ", "224 × 224"),
            ("ขนาดภาพเข้าโครงข่ายแบ่งส่วน", "256 × 256"),
            ("เกณฑ์พิกเซลขั้นต่ำของเมล็ดตับ", "100 พิกเซล"),
            ("เกณฑ์พิกเซลขั้นต่ำของเมล็ดถุงน้ำดี", "60 พิกเซล"),
            ("ขนาดตัวประกอบทำความสะอาดหน้ากากตับ", "11"),
            ("ขนาดตัวประกอบทำความสะอาดหน้ากากถุงน้ำดี", "7"),
            ("เกณฑ์ตัดพิกเซลเงาอะคูสติก", "ค่าความสว่างต่ำกว่า 8"),
        ],
        widths=[4.0, 2.5],
        align_cols=[None, WD_ALIGN_PARAGRAPH.CENTER],
    )

    h2(doc, "ก.2 โมดูลพังผืดตับ")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("มิติเวกเตอร์ลักษณะเด่นจากแกนสกัด", "512"),
            ("ขนาดตารางฝังมุมตรวจ", "4 ตำแหน่ง × 8 มิติ"),
            ("มิติที่เข้าหัวทำนาย", "520"),
            ("อัตรา dropout", "0.3"),
            ("น้ำหนักองค์ประกอบสูญเสีย ถดถอย / CORN / SWE", "1.0 / 0.3 / 0.1"),
            ("ค่า delta ของ Huber loss", "0.3"),
            ("เกณฑ์ทางคลินิก TE (kPa)", "6.0 / 7.1 / 8.7 / 10.3"),
            ("จุดตัดที่ปรับเทียบตามความชุก (สเกล log)", "1.3303 / 1.4456 / 1.5376 / 1.6352"),
            ("จุดตัดที่ปรับเทียบ เทียบเท่าเป็น kPa", "3.78 / 4.24 / 4.65 / 5.13"),
            ("จุดแบ่งระดับความเสี่ยงจากโอกาส ≥F2", "0.15 และ 0.30"),
            ("จำนวนโมเดลในคณะ", "5"),
            ("ระยะห่างเวลาสูงสุดในการจัดกลุ่มครั้งตรวจ", "300 วินาที"),
            ("จำนวนชุดการแบ่งข้อมูล", "5 ส่วน × 3 รอบ = 15 ชุด"),
            ("จำนวนครั้งตรวจระดับ F3 และ F4 ขั้นต่ำต่อชุด", "5 และ 6"),
            ("เกณฑ์ระดับที่รายงาน F4", "โอกาส F4 ≥ 0.25 หรือ (kPa ≥ 6.0 และโอกาส ≥F3 ≥ 0.25)"),
            ("เกณฑ์ระดับที่รายงาน F3", "โอกาส ≥F3 ≥ 0.40 หรือ kPa ≥ 5.5"),
            ("เกณฑ์ระดับที่รายงาน F2", "โอกาส ≥F2 ≥ 0.35 หรือ kPa ≥ 4.6"),
            ("เกณฑ์ระดับที่รายงาน F1", "โอกาส ≥F2 ≥ 0.25 หรือ kPa ≥ 4.0"),
        ],
        widths=[3.55, 2.95],
    )
    page_break(doc)

    h2(doc, "ก.3 โมดูลไขมันพอกตับ")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("แถบความลึกตื้น", "15 % ถึง 45 % ของความสูงหน้ากากตับ"),
            ("แถบความลึกลึก", "55 % ถึง 90 % ของความสูงหน้ากากตับ"),
            ("ควอนไทล์ที่ใช้ตัดปลาย", "10 และ 90"),
            ("ความสูงตับขั้นต่ำที่แบ่งแถบได้", "มากกว่า 30 พิกเซล"),
            ("อัตราการลดทอนตั้งต้นเมื่อแบ่งแถบไม่ได้", "1.05"),
            ("ชั้นคะแนนอัตราการลดทอน", "≥1.45 → +2.2 · ≥1.30 → +1.5 · ≥1.18 → +0.8"),
            ("ชั้นคะแนนความสว่างเนื้อตับ", "≥120 → +1.6 · ≥95 → +0.9 · ≥80 → +0.3"),
            ("คะแนนเมื่อพบไขมันเฉพาะที่", "+1.8"),
            ("คะแนนสูงสุดที่เป็นไปได้", "5.6"),
            ("เกณฑ์แบ่งระดับ", "≥2.6 → S3 · ≥1.6 → S2 · ≥0.8 → S1 · ต่ำกว่านั้น → S0"),
            ("ค่าความมั่นใจประจำระดับ", "S3 0.88 · S2 0.85 · S1 0.82 · S0 0.90"),
        ],
        widths=[3.1, 3.4],
    )

    h2(doc, "ก.4 โมดูลรอยโรคเฉพาะที่")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("ขนาดภาพขณะฝึก และขณะอนุมาน", "512 × 512 และ 640 × 640"),
            ("จำนวนคลาส", "7"),
            ("ค่าความมั่นใจขั้นต่ำตั้งต้น", "0.25"),
            ("เกณฑ์ IoU ของการกดทับกล่องซ้อน", "0.7"),
            ("จำนวนกล่องสูงสุดต่อภาพ", "300"),
            ("ขนาดกล่องขั้นต่ำ", "6 พิกเซล ทั้งความกว้างและความสูง"),
            ("เกณฑ์สัดส่วนการทับซ้อนกับหน้ากากตับ", "0.25"),
            ("ค่าความมั่นใจที่ติดธงเตือน", "ต่ำกว่า 0.60"),
            ("ค่าคงที่ประมาณขนาด", "0.35 มิลลิเมตรต่อพิกเซล"),
            ("จำนวนรอบการฝึก และขนาดชุดย่อย", "35 และ 32"),
            ("น้ำหนักสูญเสีย กล่อง / คลาส / DFL", "7.5 / 0.5 / 1.5"),
            ("ค่าการเพิ่มความหลากหลายที่ปรับเฉพาะ",
             "พลิกบนล่าง 0.0 · พลิกซ้ายขวา 0.5 · mosaic 0.5\n"
             "· mixup 0.1 · หมุน 10° · เลื่อน 0.10 · ย่อขยาย 0.15\n· ความสว่าง 0.3"),
        ],
        widths=[3.1, 3.4],
    )

    h2(doc, "ก.5 กฎยกระดับและโมดูลความเสี่ยงพยาธิ")
    table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("เงื่อนไขกระตุ้นกฎยกระดับ", "พบรอยโรคชนิด HCC หรือ CCA"),
            ("เงื่อนไขของโมดูลพังผืด", "ระดับความเสี่ยงต่ำกว่าสูง หรือระดับอยู่ใน F0, F1, F2"),
            ("ค่ายกพื้น kPa", "9.5"),
            ("ค่ายกพื้นโอกาสเป็น F4", "0.65"),
            ("คะแนนฐานของโมดูลพยาธิ", "0.10"),
            ("คะแนนประวัติกินปลาน้ำจืดดิบ", "+0.55"),
            ("คะแนนประวัติเคยรักษาพยาธิ", "+0.30"),
            ("คะแนนประวัติมะเร็งในครอบครัว", "+0.15"),
            ("เกณฑ์แบ่งผล", "≥0.70 → Probable · ≥0.35 → Possible · ต่ำกว่านั้น → Negative"),
        ],
        widths=[2.9, 3.6],
    )
    page_break(doc)


def appendix_b(doc):
    h1(doc, "ภาคผนวก ข  ดัชนีไฟล์และตำแหน่งบรรทัดอ้างอิง")
    body(doc, "ระบุตำแหน่งของตรรกะสำคัญทุกจุดที่อ้างถึงในเอกสาร เพื่อให้ตรวจสอบย้อนกลับได้ "
              "หมายเลขบรรทัดอ้างอิงสถานะของโครงการ ณ วันที่จัดทำเอกสาร")
    table(
        doc,
        ["หัวข้อ", "ไฟล์", "บรรทัด"],
        [
            ("ผลงานเดิม ขอบเขตและผลการวัด", "data/SmartLiva-LiverUS-SDK/README.md", "ทั้งไฟล์"),
            ("ผลงานเดิม ผลด่านอวัยวะ", "data/SmartLiva-LiverUS-SDK/runs_v3/metrics_organ.json", "—"),
            ("ผลงานเดิม ผลการแบ่งส่วน", "data/SmartLiva-LiverUS-SDK/runs_seg/metrics.json", "—"),
            ("การรับเข้าและเตรียมภาพ", "src/api/server.py", "203–218"),
            ("การสร้างหน้ากากตับของช่องทางรายโมดูล", "src/api/server.py", "221–240"),
            ("ช่องทางวิเคราะห์เต็มรูปแบบ", "src/api/server.py", "542–652"),
            ("โครงสร้างข้อมูลที่รับเข้า", "src/workflow/schemas.py", "9–36"),
            ("โครงสร้างผลลัพธ์รวม", "src/workflow/schemas.py", "105–146"),
            ("การลบข้อมูลระบุตัวตน", "src/models/segmentation/morphology.py", "18–29"),
            ("การทำความสะอาดหน้ากาก", "src/models/segmentation/morphology.py", "32–54"),
            ("การทำงานร่วมของสองโมเดลแบ่งส่วน", "src/models/segmentation/pipeline.py", "21–131"),
            ("การหักถุงน้ำดีสองรอบ", "src/models/segmentation/pipeline.py", "113–115, 126–129"),
            ("ด่านจำแนกอวัยวะและเกณฑ์ละเว้น", "src/models/gate/classify.py", "41–44, 139–155"),
            ("การรวมคำตัดสินของด่าน", "src/workflow/gatekeeper.py", "74–155"),
            ("ด่านบังคับและการหยุดการทำงาน", "src/workflow/orchestrator.py", "49–112"),
            ("ลำดับการส่งโมดูลเข้าทำงาน", "src/workflow/orchestrator.py", "126–164"),
            ("ลำดับชั้นตรวจทาน", "src/workflow/orchestrator.py", "209–215"),
            ("โครงสร้าง FibrosisNet และฟังก์ชันสูญเสีย", "src/models/fibrosis/model.py", "68–208"),
            ("เกณฑ์ระดับและการจัดกลุ่มครั้งตรวจ", "src/models/fibrosis/labels.py", "34–54, 78–110"),
            ("การเตรียม ROI", "src/models/fibrosis/preprocess.py", "33–189"),
            ("การเพิ่มความหลากหลายและ SpeckleNoise", "src/models/fibrosis/dataset.py", "39–77"),
            ("การแบ่งข้อมูลและการตรวจการรั่วไหล", "src/models/fibrosis/splits.py", "31–116"),
            ("ขั้นตอนการฝึกและการเฉลี่ยน้ำหนัก", "src/models/fibrosis/train.py", "47–72, 326–348"),
            ("การปรับเทียบเกณฑ์ตามความชุก", "src/models/fibrosis/evaluate.py", "94–113"),
            ("การอนุมานและการคำนวณ kPa", "src/models/fibrosis/infer.py", "191–274"),
            ("ระดับความเสี่ยงและอัตราที่สังเกตได้", "src/models/fibrosis/infer.py", "43–45, 78–114"),
            ("ชั้นตัดสินระดับที่รายงาน", "src/workflow/specialists/fibrosis.py", "56–86"),
            ("โมดูลไขมันพอกตับทั้งหมด", "src/workflow/specialists/fatty_liver.py", "17–192"),
            ("แถบความลึกและตัวประมาณทนทาน", "src/workflow/specialists/fatty_liver.py", "45–79"),
            ("การให้คะแนนและแบ่งระดับ", "src/workflow/specialists/fatty_liver.py", "92–144"),
            ("กลไกกักบริเวณเชิงพื้นที่", "src/workflow/specialists/lesion.py", "50–69"),
            ("การประมาณขนาดและพิกัดผลลัพธ์", "src/workflow/specialists/lesion.py", "78–98"),
            ("โมดูลความเสี่ยงพยาธิ", "src/workflow/specialists/fluke_risk.py", "22–113"),
            ("ชั้นกฎรวมผลและกฎยกระดับ", "src/workflow/verifiers/rule_engine.py", "7–53"),
            ("ตัวเรียบเรียงรายงานสำรอง", "src/workflow/verifiers/medical_reviewer.py", "17–77"),
            ("ฐานข้อมูลวงจรแพทย์ตรวจทาน", "src/database/flywheel.py", "22–162"),
            ("โครงสร้างการตรวจทานของแพทย์", "frontend/src/domain/review.ts", "18–213"),
            ("สคริปต์เตรียมข้อมูลและฝึกโมดูลรอยโรค", "tools/prep_and_train_lesion.py", "19–139"),
            ("ค่าพารามิเตอร์การฝึกที่ใช้จริง",
             "runs/lesion_train/smartliva_yolov8n_local/args.yaml", "ทั้งไฟล์"),
            ("ผลการฝึกรายรอบ", "runs/lesion_train/smartliva_yolov8n_local/results.csv", "แถวสุดท้าย"),
            ("ผลการวัดของโมดูลพังผืด", "src/models/fibrosis/reports/metrics.json", "—"),
            ("ผลการทดสอบชุดควบคุมเชิงลบ", "src/models/fibrosis/reports/verdict.json", "—"),
            ("ชุดทดสอบทางคลินิก", "tools/run_full_clinical_test_suite.py", "180–332"),
        ],
        widths=[2.35, 3.2, 0.95],
        align_cols=[None, None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    page_break(doc)


def appendix_c(doc):
    h1(doc, "ภาคผนวก ค  ประเด็นที่ต้องยืนยันก่อนยกร่างข้อถือสิทธิ")
    body(doc, "การตรวจสอบซอร์สโค้ดพบว่าเอกสารประกอบเดิมของโครงการบางส่วนบรรยายความสามารถ "
              "ที่ยังไม่ปรากฏในโค้ดที่ใช้งานจริง เอกสารฉบับนี้ไม่นำข้อความเหล่านั้นมาใช้ "
              "และรวบรวมไว้ที่นี่เพื่อให้ผู้ยกร่างคำขอพิจารณาก่อนกำหนดขอบเขตข้อถือสิทธิ "
              "เนื่องจากการกล่าวอ้างความสามารถที่ยังไม่ได้ทำจริง อาจกระทบต่อความสมบูรณ์ของคำขอในภายหลัง")

    table(
        doc,
        ["ข้อ", "ข้อกล่าวอ้างในเอกสารเดิม", "สิ่งที่พบในซอร์สโค้ด", "ข้อเสนอ"],
        [
            ("1", "ระบบชื่อคลาสรอยโรคเจ็ดชนิด\nตามที่ระบุในหัวข้อ 7.2",
             "ไฟล์น้ำหนักที่ใช้งานบันทึกรายชื่อคลาส\nไว้เป็นอีกชุดหนึ่ง ชั้นแปลงชื่อคลาส\nที่คั่นอยู่ยังไม่สอดคล้องกัน",
             "แก้ให้ตรงกันแล้ววัดผลใหม่\nก่อนอ้างอิงชนิดรอยโรค\nในข้อถือสิทธิ"),
            ("2", "โมดูลพยาธิวิเคราะห์ลักษณะพื้นผิว\nรอบท่อน้ำดี และคำนวณดัชนี\nการหนาตัวของผนังท่อ",
             "ไม่มีการคำนวณดังกล่าว โมดูลใช้ภาพ\nเพียงเพื่อหาจุดศูนย์ถ่วง คะแนนมาจาก\nประวัติผู้ป่วยล้วน",
             "ไม่ควรอ้างการวิเคราะห์ภาพ\nสำหรับโมดูลนี้ หรือพัฒนา\nส่วนดังกล่าวขึ้นก่อน"),
            ("3", "การแบ่งข้อมูลระดับผู้ป่วย\nและปลอดการรั่วไหลสำหรับ\nโมดูลรอยโรค",
             "สคริปต์ที่รันจริงใช้การสุ่มสลับ\nลำดับไฟล์แบบแบน 85 ต่อ 15\nไม่ได้แบ่งตามผู้ป่วย",
             "แบ่งใหม่ตามผู้ป่วยแล้ววัดผลซ้ำ\nหรือไม่กล่าวอ้างคุณสมบัตินี้"),
            ("4", "ค่า AUROC สำหรับการแยก\nอย่างน้อยระดับ F2 เท่ากับ 0.892",
             "ไฟล์ผลการวัดระบุ 0.7180 ± 0.0477\nเป็นค่าเฉลี่ย และ 0.7386 เป็นค่ารวม",
             "ใช้ตัวเลขจากไฟล์ผลการวัด\nเท่านั้น"),
            ("5", "โมดูลไขมันใช้ดัชนีเปรียบเทียบ\nความสว่างตับกับไต",
             "ไม่มีการแบ่งส่วนไตในระบบเลย\nโมดูลใช้เพียงอัตราการลดทอน\nและความสว่างเนื้อตับ",
             "ไม่ควรอ้างดัชนีดังกล่าว"),
            ("6", "คณะโมเดลพังผืดประกอบด้วย\nสถาปัตยกรรมสองแบบผสมกัน",
             "ไฟล์น้ำหนักที่ใช้งานเป็น ResNet-18\nทั้งห้าตัว ไม่มีสถาปัตยกรรมอื่น\nอยู่บนดิสก์",
             "ระบุว่าเป็น ResNet-18\nอย่างเดียว"),
            ("7", "การระบุตำแหน่งรอยโรค\nเป็นเซกเมนต์ตับ",
             "ไม่มีการตรวจหาจุดสังเกตทางกายวิภาค\nระบบรายงานเพียงกล่องพิกัด\nปรับมาตรฐาน",
             "ไม่ควรอ้างการแบ่งเซกเมนต์"),
            ("8", "ด่านฟิสิกส์ของภาพทำงาน\nในทุกเส้นทางการเรียกใช้",
             "ในเส้นทางที่เรียกผ่านบริการเว็บ\nด่านฟิสิกส์คืนสถานะไม่พร้อมใช้งาน\nและถูกข้ามไป",
             "แก้การเชื่อมต่อ หรือจำกัด\nการอ้างอิงให้ตรงกับเส้นทาง\nที่ทำงานจริง"),
            ("9", "โมดูลความเสี่ยงพยาธิให้ผล\nระดับสูงได้",
             "ในช่องทางวิเคราะห์เต็มรูปแบบ\nคะแนนสูงสุดคือ 0.65 ซึ่งไม่ถึงเกณฑ์\n0.70 จึงไม่ถึงระดับสูง",
             "เพิ่มเขตข้อมูลประวัติที่ขาด\nลงในโครงสร้างข้อมูล"),
            ("10", "กฎยกระดับพื้นหลังตับแข็ง\nใช้กับมะเร็งตับชนิดปฐมภูมิ",
             "โค้ดกระตุ้นกฎเดียวกันเมื่อพบ\nมะเร็งท่อน้ำดีด้วย ซึ่งไม่ใช่ตัวบ่งชี้\nภาวะตับแข็งในทางคลินิก",
             "ยืนยันเจตนา แล้วแก้โค้ด\nหรือแก้คำบรรยายให้ตรงกัน"),
        ],
        widths=[0.45, 1.95, 2.35, 1.75],
        align_cols=[WD_ALIGN_PARAGRAPH.CENTER, None, None, None],
    )

    body(doc, "นอกจากสิบข้อข้างต้น ยังมีประเด็นเชิงเทคนิคที่ควรพิจารณาแก้ไข แม้ไม่กระทบต่อสาระ "
              "ของสิ่งที่ขอความคุ้มครองโดยตรง ได้แก่ ขนาดภาพขณะอนุมานของโมดูลรอยโรคต่างจากขณะฝึก "
              "ค่าลำดับระดับพังผืดในผลลัพธ์ไม่ถูกปรับตามเมื่อกฎยกระดับทำงาน "
              "และเกณฑ์พิกเซลขั้นต่ำของถุงน้ำดีที่ใช้ในชั้นแบ่งส่วนกับชั้นรายงานผลไม่ตรงกัน")

    rule(doc)
    body(doc, "จบเอกสาร")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    setup_page(doc)
    add_page_number_footer(doc)

    cover(doc)
    contents(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)

    part_divider(
        doc,
        "ภาค ก",
        "โมเดลพังผืดตับและไขมันพอกตับ",
        [
            "4.  โมเดลจัดระดับพังผืดตับ F0–F4 (FibrosisNet Ensemble)",
            "5.  โมดูลประเมินไขมันพอกตับ S0–S3",
            "6.  จุดร่วมภายในภาค ก",
        ],
    )
    section_4(doc)
    section_5(doc)
    section_6(doc)

    part_divider(
        doc,
        "ภาค ข",
        "โมเดลรอยโรค มะเร็ง และพยาธิใบไม้ตับ",
        [
            "7.  ตัวตรวจจับรอยโรคเฉพาะที่ในเนื้อตับ",
            "8.  กฎยกระดับเมื่อพบรอยโรคมะเร็ง และการส่งตรวจยืนยัน",
            "9.  โมดูลประเมินความเสี่ยงพยาธิใบไม้ตับและมะเร็งท่อน้ำดี",
            "10.  จุดร่วมภายในภาค ข",
        ],
    )
    section_7(doc)
    section_8(doc)
    section_9(doc)
    section_10(doc)

    part_divider(
        doc,
        "ส่วนที่ 3",
        "จุดร่วมของแต่ละโมเดล",
        ["11.  จุดร่วมของทั้งสี่โมดูล"],
    )
    section_11(doc)

    appendix_a(doc)
    appendix_b(doc)
    appendix_c(doc)

    out_dir = Path(__file__).resolve().parent.parent / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "SmartLiva_IP_Technical_Disclosure_TH.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    main()
